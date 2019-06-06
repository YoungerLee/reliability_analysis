from Ui_reliability_main import Ui_MainWindow
from output_report_dialog import OutputReportDialog
from table_list import *
from meta.report import Report
import reliability_analysis.function as rf
import reliability_analysis.estimate as estimate
from PyQt5.QtCore import pyqtSlot, Qt, QTime, QEventLoop, QCoreApplication
from PyQt5.QtChart import QLineSeries, QScatterSeries, QChart, QBarSet, QBarSeries, QBarCategoryAxis, QValueAxis
from PyQt5.QtWidgets import QMainWindow, QApplication, QFileDialog, QMessageBox
from PyQt5.QtGui import QFont, QPainter, QColor
from PyQt5 import QtSql
from PyQt5.QtSql import QSqlQuery
import numpy as np
from enum import Enum, unique
from openpyxl import Workbook
import os
import operator
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

@unique
class Result(Enum):
    lamda = 1
    lamda_value = 2
    beta = 3
    beta_value = 4
    mtbf = 5
    mtbf_value = 6
    curr_time = 7
    curr_time_value = 8
    pdf = 9
    pdf_value = 10
    cdf = 11
    cdf_value = 12
    relia = 13
    relia_value = 14
    fali = 15
    fali_value = 16

@unique
class Fault(Enum):
    pattern = 0
    pattern_value = 1
    position = 2
    position_value = 3
    reason = 4
    reason_value = 5
    root = 6
    root_value = 7

[RUN_TIME, ENV_TEMP, KNI_TEMP, RPA] = range(4)
[PATT, POSI, REASON, ROOT] = range(4)

def delay(msec):
    dieTime = QTime.currentTime().addMSecs(msec)
    while QTime.currentTime() < dieTime:
        QCoreApplication.processEvents(QEventLoop.AllEvents, 100)


class MainWindow(QMainWindow):
    def __init__(self, parent=None):
        super(MainWindow, self).__init__(parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.lambda_hat = -1
        self.beta_hat = -1
        self.lambda_map = -1
        self.beta_map = -1
        self.lambda_hat_2 = -1
        self.beta_hat_2 = -1
        self.lambda_map_2 = -1
        self.beta_map_2 = -1
        self.data_table = dict()
        self.ui.stackedWidget.setCurrentIndex(0)
        self.ui.listWidget.itemClicked.connect(self.on_listWidget_itemClicked)
        self.ui.comboBox.currentIndexChanged.connect(self.on_comboBox_currentIndexChanged)
        self.ui.comboBox_2.currentIndexChanged.connect(self.on_comboBox_2_currentIndexChanged)
        self.ui.devComboBox.currentIndexChanged.connect(self.on_devComboBox_currentIndexChanged)
        self.ui.devComboBox_2.currentIndexChanged.connect(self.on_devComboBox_2_currentIndexChanged)
        self.ui.axisComboBox.currentIndexChanged.connect(self.on_axisComboBox_currentIndexChanged)
        self.ui.outputReportAction.triggered.connect(self.outputReport)
        self.initScatterCharts()
        self.initLinePlot()
        self.initBarPlot()
        self.initRawDataCharts()
        self.initDataBase()
        self.initDevInfo()

    def initDataBase(self):
        database = QtSql.QSqlDatabase.addDatabase('QSQLITE')
        database.setDatabaseName('./db/reliability.db')
        database.open()
        self.query = QSqlQuery()
        if not self.is_table_existed('device'):
            sql_str = 'create table device (id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE NOT NULL, ' \
                      'num varchar(50), name varchar(50))'
            self.query.prepare(sql_str)
            if not self.query.exec_():
                print(self.query.lastError().text())
            else:
                print('create a device table')

        if not self.is_table_existed('record'):
            sql_str = "create table record (id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE NOT NULL, " \
                      "run_time double, env_temp double, kni_temp double, rpa double, axis variable(5), " \
                      "dev_id int, record_time timestamp NOT NULL DEFAULT(datetime('now','localtime')), " \
                      "foreign key(dev_id) references device(id))"
            self.query.prepare(sql_str)
            if not self.query.exec_():
                print(self.query.lastError().text())
            else:
                print('create a record table')
        if not self.is_table_existed('maintain'):
            sql_str = "create table maintain (id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE NOT NULL, " \
                      "maintain_date date, person varchar(50), maintain_time double, " \
                      "dev_id int, record_time timestamp NOT NULL DEFAULT(datetime('now','localtime')), " \
                      "foreign key(dev_id) references device(id))"
            self.query.prepare(sql_str)
            if not self.query.exec_():
                print(self.query.lastError().text())
            else:
                print('create a maintain table')
        if not self.is_table_existed('breakdown'):
            sql_str = "create table breakdown (id INTEGER PRIMARY KEY AUTOINCREMENT UNIQUE NOT NULL, " \
                      "pattern varchar(50), position varchar(50), reason varchar(50), root varchar(50), " \
                      "status bit, dev_id int, record_time timestamp NOT NULL DEFAULT(datetime('now','localtime')), " \
                      "foreign key(dev_id) references device(id))"
            self.query.prepare(sql_str)
            if not self.query.exec_():
                print(self.query.lastError().text())
            else:
                print('create a breakdown table')

    def is_table_existed(self, table_name):
        existed = False
        sql_str = "select count(*) from sqlite_master where type='table' and name='{0}'".format(table_name)
        self.query.exec(sql_str)
        if self.query.next():
            if int(self.query.value(0)) == 0:  # 表不存在
                existed = False
            else:  # 表存在
                existed = True
        return existed

    def initDevInfo(self):
        self.ui.devComboBox.clear()
        self.ui.devComboBox.addItem('选择设备')
        self.ui.devComboBox_2.clear()
        self.ui.devComboBox_2.addItem('选择设备')
        query_sql = 'select name from device'
        self.query.prepare(query_sql)
        if not self.query.exec_():
            print(self.query.lastError().text())
        else:
            while self.query.next():
                name = self.query.value(0)
                self.ui.devComboBox.addItem(name)
                self.ui.devComboBox_2.addItem(name)

    def initScatterCharts(self):
        '''
        初始化散点图
        :return: None
        '''
        ##---可靠度分析---
        self.m_scatter_chart = QChart()
        self.m_scatter_series = QScatterSeries(self.m_scatter_chart)
        self.m_scatter_series_2 = QScatterSeries(self.m_scatter_chart)  # dev_2
        self.m_scatter_series.setName('设备1')
        self.m_scatter_series_2.setName('设备2')
        self.m_scatter_chart.addSeries(self.m_scatter_series)
        self.m_scatter_chart.addSeries(self.m_scatter_series_2)
        self.m_scatter_chart.setTitle("故障间隔时间散点图")  # 设置图题
        self.m_scatter_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_scatter_chart.createDefaultAxes()  # 创建默认轴
        self.m_scatter_chart.axisX().setTitleText('故障间隔时间/h')  # 设置横坐标标题
        self.m_scatter_chart.axisX().setRange(0, 1)
        self.m_scatter_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_scatter_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_scatter_chart.axisY().setTitleText('概率')  # 设置纵坐标标题
        self.m_scatter_chart.axisY().setRange(0, 1)
        self.m_scatter_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_scatter_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.scatterView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.scatterView.setChart(self.m_scatter_chart)
        self.ui.scatterView.show()
        ##--维修性分析--
        self.m_scatter_chart_maintain = QChart()
        self.m_scatter_series_maintain = QScatterSeries()
        self.m_scatter_series_maintain_2 = QScatterSeries()  # dev_2
        self.m_scatter_series_maintain.setName('设备1')
        self.m_scatter_series_maintain_2.setName('设备2')
        self.m_scatter_chart_maintain.addSeries(self.m_scatter_series_maintain)
        self.m_scatter_chart_maintain.addSeries(self.m_scatter_series_maintain_2)
        self.m_scatter_chart_maintain.setTitle("维修时间散点图")  # 设置图题
        self.m_scatter_chart_maintain.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_scatter_chart_maintain.createDefaultAxes()  # 创建默认轴
        self.m_scatter_chart_maintain.axisX().setTitleText('维修时间/h')  # 设置横坐标标题
        self.m_scatter_chart_maintain.axisX().setRange(0, 1)
        self.m_scatter_chart_maintain.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_scatter_chart_maintain.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_scatter_chart_maintain.axisY().setTitleText('概率')  # 设置纵坐标标题
        self.m_scatter_chart_maintain.axisY().setRange(0, 1)
        self.m_scatter_chart_maintain.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_scatter_chart_maintain.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.scatterView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.scatterView_2.setChart(self.m_scatter_chart_maintain)
        self.ui.scatterView_2.show()

    def initLinePlot(self):
        '''
        初始化曲线图
        :return:
        '''
        # ---可靠度分析---
        # 初始化概率密度图
        self.m_pdf_chart = QChart()
        # init series
        self.m_pdf_series_ls = QLineSeries()
        self.m_pdf_series_ls_2 = QLineSeries()  # dev_2
        self.m_pdf_series_map = QLineSeries()
        self.m_pdf_series_map_2 = QLineSeries()  # dev_2
        # connect slot
        self.m_pdf_series_ls.hovered.connect(self.on_m_pdf_series_hovered)  # 响应显示曲线数值事件
        self.m_pdf_series_ls_2.hovered.connect(self.on_m_pdf_series_hovered)
        self.m_pdf_series_map.hovered.connect(self.on_m_pdf_series_hovered)  # 响应显示曲线数值事件
        self.m_pdf_series_map_2.hovered.connect(self.on_m_pdf_series_hovered)
        # add series
        self.m_pdf_series_ls.setName('设备1-LS')
        self.m_pdf_series_map.setName('设备1-MAP')
        self.m_pdf_series_ls_2.setName('设备2-LS')
        self.m_pdf_series_map_2.setName('设备2-MAP')
        self.m_pdf_chart.addSeries(self.m_pdf_series_ls)
        self.m_pdf_chart.addSeries(self.m_pdf_series_map)
        self.m_pdf_chart.addSeries(self.m_pdf_series_ls_2)
        self.m_pdf_chart.addSeries(self.m_pdf_series_map_2)
        # set title and style
        self.m_pdf_chart.setTitle('威布尔分布概率密度图')  # 设置图题
        self.m_pdf_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_pdf_chart.createDefaultAxes()  # 创建默认轴
        self.m_pdf_chart.axisX().setTitleText('故障间隔时间/h')  # 设置横坐标标题
        self.m_pdf_chart.axisX().setRange(0, 1)
        self.m_pdf_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_pdf_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_pdf_chart.axisY().setTitleText('概率')  # 设置纵坐标标题
        self.m_pdf_chart.axisY().setRange(0, 10e-4)
        self.m_pdf_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_pdf_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.pdfView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.pdfView.setChart(self.m_pdf_chart)
        self.ui.pdfView.show()
        # 初始化累积分布图
        self.m_cdf_chart = QChart()
        # init series
        self.m_cdf_series_ls = QLineSeries()
        self.m_cdf_series_map = QLineSeries()
        self.m_cdf_series_ls_2 = QLineSeries()  # dev_2
        self.m_cdf_series_map_2 = QLineSeries()  # dev_2
        # connect slot
        self.m_cdf_series_ls.hovered.connect(self.on_m_cdf_series_hovered)  # 响应显示曲线数值事件
        self.m_cdf_series_map.hovered.connect(self.on_m_cdf_series_hovered)  # 响应显示曲线数值事件
        self.m_cdf_series_ls_2.hovered.connect(self.on_m_cdf_series_hovered)
        self.m_cdf_series_map_2.hovered.connect(self.on_m_cdf_series_hovered)
        # add series
        self.m_cdf_series_ls.setName('设备1-LS')
        self.m_cdf_series_map.setName('设备1-MAP')
        self.m_cdf_series_ls_2.setName('设备2-LS')
        self.m_cdf_series_map_2.setName('设备2-MAP')
        self.m_cdf_chart.addSeries(self.m_cdf_series_ls)
        self.m_cdf_chart.addSeries(self.m_cdf_series_map)
        self.m_cdf_chart.addSeries(self.m_cdf_series_ls_2)
        self.m_cdf_chart.addSeries(self.m_cdf_series_map_2)
        # set title and style
        self.m_cdf_chart.setTitle("威布尔分布累积分布图")  # 设置图题
        self.m_cdf_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_cdf_chart.createDefaultAxes()  # 创建默认轴
        self.m_cdf_chart.axisX().setTitleText('故障间隔时间/h')  # 设置横坐标标题
        self.m_cdf_chart.axisX().setRange(0, 1)
        self.m_cdf_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_cdf_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_cdf_chart.axisY().setTitleText('概率')  # 设置纵坐标标题
        self.m_cdf_chart.axisY().setRange(0, 1)
        self.m_cdf_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_cdf_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.cdfView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.cdfView.setChart(self.m_cdf_chart)
        self.ui.cdfView.show()
        # 初始化可靠度曲线图
        self.m_relia_chart = QChart()
        # init series
        self.m_relia_series_ls = QLineSeries()
        self.m_relia_series_map = QLineSeries()
        self.m_relia_series_ls_2 = QLineSeries()  # dev_2
        self.m_relia_series_map_2 = QLineSeries()
        # connect slot
        self.m_relia_series_ls.hovered.connect(self.on_m_relia_series_hovered)  # 响应显示曲线数值事件
        self.m_relia_series_map.hovered.connect(self.on_m_relia_series_hovered)  # 响应显示曲线数值事件
        self.m_relia_series_ls_2.hovered.connect(self.on_m_relia_series_hovered)
        self.m_relia_series_map_2.hovered.connect(self.on_m_relia_series_hovered)
        # add series
        self.m_relia_series_ls.setName('设备1-LS')
        self.m_relia_series_map.setName('设备1-MAP')
        self.m_relia_series_ls_2.setName('设备2-LS')
        self.m_relia_series_map_2.setName('设备2-MAP')
        self.m_relia_chart.addSeries(self.m_relia_series_ls)
        self.m_relia_chart.addSeries(self.m_relia_series_map)
        self.m_relia_chart.addSeries(self.m_relia_series_ls_2)
        self.m_relia_chart.addSeries(self.m_relia_series_map_2)
        # set title and style
        self.m_relia_chart.setTitle("可靠度曲线图")  # 设置图题
        self.m_relia_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_relia_chart.createDefaultAxes()  # 创建默认轴
        self.m_relia_chart.axisX().setTitleText('故障间隔时间/h')  # 设置横坐标标题
        self.m_relia_chart.axisX().setRange(0, 1)
        self.m_relia_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_relia_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_relia_chart.axisY().setTitleText('可靠度')  # 设置纵坐标标题
        self.m_relia_chart.axisY().setRange(0, 1)
        self.m_relia_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_relia_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.reliaView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.reliaView.setChart(self.m_relia_chart)
        self.ui.reliaView.show()
        # 初始化失效率曲线图
        self.m_fail_chart = QChart()
        # init series
        self.m_fail_series_ls = QLineSeries()
        self.m_fail_series_map = QLineSeries()
        self.m_fail_series_ls_2 = QLineSeries()  # dev_2
        self.m_fail_series_map_2 = QLineSeries()
        # connect slot
        self.m_fail_series_ls.hovered.connect(self.on_m_fali_series_hovered)  # 响应显示曲线数值事件
        self.m_fail_series_map.hovered.connect(self.on_m_fali_series_hovered)  # 响应显示曲线数值事件
        self.m_fail_series_ls_2.hovered.connect(self.on_m_fali_series_hovered)
        self.m_fail_series_map_2.hovered.connect(self.on_m_fali_series_hovered)
        # add series
        self.m_fail_series_ls.setName('设备1-LS')
        self.m_fail_series_map.setName('设备1-MAP')
        self.m_fail_series_ls_2.setName('设备2-LS')
        self.m_fail_series_map_2.setName('设备2-MAP')
        self.m_fail_chart.addSeries(self.m_fail_series_ls)
        self.m_fail_chart.addSeries(self.m_fail_series_map)
        self.m_fail_chart.addSeries(self.m_fail_series_ls_2)
        self.m_fail_chart.addSeries(self.m_fail_series_map_2)
        self.m_fail_chart.setTitle("失效率曲线图")  # 设置图题
        self.m_fail_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_fail_chart.createDefaultAxes()  # 创建默认轴
        self.m_fail_chart.axisX().setTitleText('故障间隔时间/h')  # 设置横坐标标题
        self.m_fail_chart.axisX().setRange(0, 1)
        self.m_fail_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_fail_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_fail_chart.axisY().setTitleText('失效率')  # 设置纵坐标标题
        self.m_fail_chart.axisY().setRange(0, 10e-3)
        self.m_fail_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_fail_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.failView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.failView.setChart(self.m_fail_chart)
        self.ui.failView.show()
        ##---维修性分析---
        # 初始化概率密度图
        self.m_pdf_chart_maintain = QChart()
        # init series
        self.m_pdf_series_maintain_ls = QLineSeries()
        self.m_pdf_series_maintain_ls_2 = QLineSeries()  # dev_2
        self.m_pdf_series_maintain_map = QLineSeries()
        self.m_pdf_series_maintain_map_2 = QLineSeries()
        # connect slot
        self.m_pdf_series_maintain_ls.hovered.connect(self.on_m_pdf_series_maintain_hovered)  # 响应显示曲线数值事件
        self.m_pdf_series_maintain_ls_2.hovered.connect(self.on_m_pdf_series_maintain_hovered)
        self.m_pdf_series_maintain_map.hovered.connect(self.on_m_pdf_series_maintain_hovered)
        self.m_pdf_series_maintain_map_2.hovered.connect(self.on_m_pdf_series_maintain_hovered)
        # add series
        self.m_pdf_series_maintain_ls.setName('设备1-LS')
        self.m_pdf_series_maintain_ls_2.setName('设备2-LS')
        self.m_pdf_series_maintain_map.setName('设备1-MAP')
        self.m_pdf_series_maintain_map_2.setName('设备2-MAP')
        self.m_pdf_chart_maintain.addSeries(self.m_pdf_series_maintain_ls)
        self.m_pdf_chart_maintain.addSeries(self.m_pdf_series_maintain_ls_2)
        self.m_pdf_chart_maintain.addSeries(self.m_pdf_series_maintain_map)
        self.m_pdf_chart_maintain.addSeries(self.m_pdf_series_maintain_map_2)
        # set style
        self.m_pdf_chart_maintain.setTitle('威布尔分布概率密度图')  # 设置图题
        self.m_pdf_chart_maintain.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_pdf_chart_maintain.createDefaultAxes()  # 创建默认轴
        self.m_pdf_chart_maintain.axisX().setTitleText('维修时间/h')  # 设置横坐标标题
        self.m_pdf_chart_maintain.axisX().setRange(0, 1)
        self.m_pdf_chart_maintain.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_pdf_chart_maintain.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_pdf_chart_maintain.axisY().setTitleText('概率')  # 设置纵坐标标题
        self.m_pdf_chart_maintain.axisY().setRange(0, 10e-4)
        self.m_pdf_chart_maintain.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_pdf_chart_maintain.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.pdfView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.pdfView_2.setChart(self.m_pdf_chart_maintain)
        self.ui.pdfView_2.show()
        # 初始化累积分布图
        self.m_cdf_chart_maintain = QChart()
        # init series
        self.m_cdf_series_maintain_ls = QLineSeries()
        self.m_cdf_series_maintain_ls_2 = QLineSeries()
        self.m_cdf_series_maintain_map = QLineSeries()
        self.m_cdf_series_maintain_map_2 = QLineSeries()
        # connect slot
        self.m_cdf_series_maintain_ls.hovered.connect(self.on_m_cdf_series_maintain_hovered)  # 响应显示曲线数值事件
        self.m_cdf_series_maintain_ls_2.hovered.connect(self.on_m_cdf_series_maintain_hovered)
        self.m_cdf_series_maintain_map.hovered.connect(self.on_m_cdf_series_maintain_hovered)
        self.m_cdf_series_maintain_map_2.hovered.connect(self.on_m_cdf_series_maintain_hovered)
        # add series
        self.m_cdf_series_maintain_ls.setName('设备1-LS')
        self.m_cdf_series_maintain_ls_2.setName('设备2-LS')
        self.m_cdf_series_maintain_map.setName('设备1-MAP')
        self.m_cdf_series_maintain_map_2.setName('设备2-MAP')
        self.m_cdf_chart_maintain.addSeries(self.m_cdf_series_maintain_ls)
        self.m_cdf_chart_maintain.addSeries(self.m_cdf_series_maintain_ls_2)
        self.m_cdf_chart_maintain.addSeries(self.m_cdf_series_maintain_map)
        self.m_cdf_chart_maintain.addSeries(self.m_cdf_series_maintain_map_2)
        # set style
        self.m_cdf_chart_maintain.setTitle("威布尔分布累积分布图")  # 设置图题
        self.m_cdf_chart_maintain.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_cdf_chart_maintain.createDefaultAxes()  # 创建默认轴
        self.m_cdf_chart_maintain.axisX().setTitleText('维修时间/h')  # 设置横坐标标题
        self.m_cdf_chart_maintain.axisX().setRange(0, 1)
        self.m_cdf_chart_maintain.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_cdf_chart_maintain.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_cdf_chart_maintain.axisY().setTitleText('概率')  # 设置纵坐标标题
        self.m_cdf_chart_maintain.axisY().setRange(0, 1)
        self.m_cdf_chart_maintain.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_cdf_chart_maintain.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.cdfView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.cdfView_2.setChart(self.m_cdf_chart_maintain)
        self.ui.cdfView_2.show()
        # 初始化维修度曲线图
        self.m_relia_chart_maintain = QChart()
        # init series
        self.m_relia_series_maintain_ls = QLineSeries()
        self.m_relia_series_maintain_ls_2 = QLineSeries()
        self.m_relia_series_maintain_map = QLineSeries()
        self.m_relia_series_maintain_map_2 = QLineSeries()
        # connect slot
        self.m_relia_series_maintain_ls.hovered.connect(self.on_m_relia_series_maintain_hovered)  # 响应显示曲线数值事件
        self.m_relia_series_maintain_ls_2.hovered.connect(self.on_m_relia_series_maintain_hovered)
        self.m_relia_series_maintain_map.hovered.connect(self.on_m_relia_series_maintain_hovered)
        self.m_relia_series_maintain_map_2.hovered.connect(self.on_m_relia_series_maintain_hovered)
        # add series
        self.m_relia_series_maintain_ls.setName('设备1-LS')
        self.m_relia_series_maintain_ls_2.setName('设备2-LS')
        self.m_relia_series_maintain_map.setName('设备1-MAP')
        self.m_relia_series_maintain_map_2.setName('设备2-MAP')
        self.m_relia_chart_maintain.addSeries(self.m_relia_series_maintain_ls)
        self.m_relia_chart_maintain.addSeries(self.m_relia_series_maintain_ls_2)
        self.m_relia_chart_maintain.addSeries(self.m_relia_series_maintain_map)
        self.m_relia_chart_maintain.addSeries(self.m_relia_series_maintain_map_2)
        # set style
        self.m_relia_chart_maintain.setTitle("修复率曲线图")  # 设置图题
        self.m_relia_chart_maintain.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_relia_chart_maintain.createDefaultAxes()  # 创建默认轴
        self.m_relia_chart_maintain.axisX().setTitleText('维修时间/h')  # 设置横坐标标题
        self.m_relia_chart_maintain.axisX().setRange(0, 1)
        self.m_relia_chart_maintain.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_relia_chart_maintain.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_relia_chart_maintain.axisY().setTitleText('修复率')  # 设置纵坐标标题
        self.m_relia_chart_maintain.axisY().setRange(0, 1)
        self.m_relia_chart_maintain.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_relia_chart_maintain.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.reliaView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.reliaView_2.setChart(self.m_relia_chart_maintain)
        self.ui.reliaView_2.show()
        # 初始化失效率曲线图
        self.m_fail_chart_maintain = QChart()
        # init series
        self.m_fail_series_maintain_ls = QLineSeries()
        self.m_fail_series_maintain_ls_2 = QLineSeries()
        self.m_fail_series_maintain_map = QLineSeries()
        self.m_fail_series_maintain_map_2 = QLineSeries()
        # connect series
        self.m_fail_series_maintain_ls.hovered.connect(self.on_m_fali_series_maintain_hovered)  # 响应显示曲线数值事件
        self.m_fail_series_maintain_ls_2.hovered.connect(self.on_m_fali_series_maintain_hovered)
        self.m_fail_series_maintain_map.hovered.connect(self.on_m_fali_series_maintain_hovered)
        self.m_fail_series_maintain_map_2.hovered.connect(self.on_m_fali_series_maintain_hovered)
        # add series
        self.m_fail_series_maintain_ls.setName('设备1-LS')
        self.m_fail_series_maintain_ls_2.setName('设备2-LS')
        self.m_fail_series_maintain_map.setName('设备1-MAP')
        self.m_fail_series_maintain_map_2.setName('设备2-MAP')
        self.m_fail_chart_maintain.addSeries(self.m_fail_series_maintain_ls)
        self.m_fail_chart_maintain.addSeries(self.m_fail_series_maintain_ls_2)
        self.m_fail_chart_maintain.addSeries(self.m_fail_series_maintain_map)
        self.m_fail_chart_maintain.addSeries(self.m_fail_series_maintain_map_2)
        self.m_fail_chart_maintain.setTitle("失效率曲线图")  # 设置图题
        self.m_fail_chart_maintain.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_fail_chart_maintain.createDefaultAxes()  # 创建默认轴
        self.m_fail_chart_maintain.axisX().setTitleText('故障间隔时间/h')  # 设置横坐标标题
        self.m_fail_chart_maintain.axisX().setRange(0, 1)
        self.m_fail_chart_maintain.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_fail_chart_maintain.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_fail_chart_maintain.axisY().setTitleText('失效率')  # 设置纵坐标标题
        self.m_fail_chart_maintain.axisY().setRange(0, 10e-4)
        self.m_fail_chart_maintain.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_fail_chart_maintain.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.ui.failView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.failView_2.setChart(self.m_fail_chart_maintain)
        self.ui.failView_2.show()

    def initRawDataCharts(self):
        ## 环境温度
        self.m_env_temp_chart = QChart()
        # init series
        self.m_env_temp_series = QLineSeries()
        self.m_env_temp_series.setName('设备1')
        self.m_env_temp_series.setColor(QColor(0, 0, 255))
        self.m_env_temp_scatter = QScatterSeries()
        self.m_env_temp_scatter.setMarkerSize(10)
        self.m_env_temp_scatter.setColor(QColor(0, 0, 255))
        self.m_env_temp_series_2 = QLineSeries()  # dev_2
        self.m_env_temp_series_2.setName('设备2')
        self.m_env_temp_series_2.setColor(QColor(0, 255, 0))
        self.m_env_temp_scatter_2 = QScatterSeries()
        self.m_env_temp_scatter_2.setMarkerSize(10)
        self.m_env_temp_scatter_2.setColor(QColor(0, 255, 0))
        self.m_env_temp_chart.addSeries(self.m_env_temp_series)
        self.m_env_temp_chart.addSeries(self.m_env_temp_scatter)
        self.m_env_temp_chart.addSeries(self.m_env_temp_series_2)
        self.m_env_temp_chart.addSeries(self.m_env_temp_scatter_2)
        # set style
        self.m_env_temp_chart.setTitle("环境温度——时间折线图")  # 设置图题
        self.m_env_temp_chart.setTitleFont(QFont('SimHei', 15))  # 设置图题字体的类型和大小
        self.m_env_temp_chart.createDefaultAxes()  # 创建默认轴
        self.m_env_temp_chart.axisX().setTitleText('运行时间/h')  # 设置横坐标标题
        self.m_env_temp_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_env_temp_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_env_temp_chart.axisY().setTitleText('温度/℃')  # 设置纵坐标标题
        self.m_env_temp_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_env_temp_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_env_temp_chart.legend().hide()
        self.ui.tempView_1.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.tempView_1.setChart(self.m_env_temp_chart)
        self.ui.tempView_1.show()
        ## 刀头温度
        self.m_kni_temp_chart = QChart()
        # init series
        self.m_kni_temp_series = QLineSeries()
        self.m_kni_temp_series.setName('设备1')
        self.m_kni_temp_series.setColor(QColor(0, 0, 255))
        self.m_kni_temp_scatter = QScatterSeries()
        self.m_kni_temp_scatter.setMarkerSize(10)
        self.m_kni_temp_scatter.setColor(QColor(0, 0, 255))
        self.m_kni_temp_series_2 = QLineSeries()
        self.m_kni_temp_series_2.setName('设备2')
        self.m_kni_temp_series_2.setColor(QColor(0, 255, 0))
        self.m_kni_temp_scatter_2 = QScatterSeries()
        self.m_kni_temp_scatter_2.setMarkerSize(10)
        self.m_kni_temp_scatter_2.setColor(QColor(0, 255, 0))
        self.m_kni_temp_chart.addSeries(self.m_kni_temp_series)
        self.m_kni_temp_chart.addSeries(self.m_kni_temp_scatter)
        self.m_kni_temp_chart.addSeries(self.m_kni_temp_series_2)
        self.m_kni_temp_chart.addSeries(self.m_kni_temp_scatter_2)
        # set style
        self.m_kni_temp_chart.setTitle("刀头温度——时间折线图")  # 设置图题
        self.m_kni_temp_chart.setTitleFont(QFont('SimHei', 15))  # 设置图题字体的类型和大小
        self.m_kni_temp_chart.createDefaultAxes()  # 创建默认轴
        self.m_kni_temp_chart.axisX().setTitleText('运行时间/h')  # 设置横坐标标题
        self.m_kni_temp_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_kni_temp_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_kni_temp_chart.axisY().setTitleText('温度/℃')  # 设置纵坐标标题
        self.m_kni_temp_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_kni_temp_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_kni_temp_chart.legend().hide()
        self.ui.tempView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.tempView_2.setChart(self.m_kni_temp_chart)
        self.ui.tempView_2.show()
        ## 重复定位精度
        self.m_rpa_chart = QChart()
        # init series
        self.m_rpa_series = QLineSeries()
        self.m_rpa_series.setName('设备1')
        self.m_rpa_series.setColor(QColor(0, 0, 255))
        self.m_rpa_scatter = QScatterSeries()
        self.m_rpa_scatter.setMarkerSize(10)
        self.m_rpa_scatter.setColor(QColor(0, 0, 255))
        self.m_rpa_series_2 = QLineSeries()
        self.m_rpa_series_2.setName('设备2')
        self.m_rpa_series_2.setColor(QColor(0, 255, 0))
        self.m_rpa_scatter_2 = QScatterSeries()
        self.m_rpa_scatter_2.setMarkerSize(10)
        self.m_rpa_scatter_2.setColor(QColor(0, 255, 0))
        self.m_rpa_chart.addSeries(self.m_rpa_scatter)
        self.m_rpa_chart.addSeries(self.m_rpa_series)
        self.m_rpa_chart.addSeries(self.m_rpa_scatter_2)
        self.m_rpa_chart.addSeries(self.m_rpa_series_2)
        # set style
        self.m_rpa_chart.setTitle("重复定位精度——时间折线图")  # 设置图题
        self.m_rpa_chart.setTitleFont(QFont('SimHei', 15))  # 设置图题字体的类型和大小
        self.m_rpa_chart.createDefaultAxes()  # 创建默认轴
        self.m_rpa_chart.axisX().setTitleText('运行时间/h')  # 设置横坐标标题
        self.m_rpa_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_rpa_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_rpa_chart.axisY().setTitleText('重复定位精度/μm')  # 设置纵坐标标题
        self.m_rpa_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_rpa_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_rpa_chart.legend().hide()
        self.ui.rpaView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.rpaView.setChart(self.m_rpa_chart)
        self.ui.rpaView.show()
        ## 温差
        self.m_stra_chart = QChart()
        # init series
        self.m_stra_series = QLineSeries()
        self.m_stra_series.setName('设备1')
        self.m_stra_series.setColor(QColor(0, 0, 255))
        self.m_stra_scatter = QScatterSeries()
        self.m_stra_scatter.setMarkerSize(10)
        self.m_stra_scatter.setColor(QColor(0, 0, 255))
        self.m_stra_series_2 = QLineSeries()
        self.m_stra_series_2.setName('设备2')
        self.m_stra_series_2.setColor(QColor(0, 255, 0))
        self.m_stra_scatter_2 = QScatterSeries()
        self.m_stra_scatter_2.setMarkerSize(10)
        self.m_stra_scatter_2.setColor(QColor(0, 255, 0))
        self.m_stra_chart.addSeries(self.m_stra_scatter)
        self.m_stra_chart.addSeries(self.m_stra_series)
        self.m_stra_chart.addSeries(self.m_stra_scatter_2)
        self.m_stra_chart.addSeries(self.m_stra_series_2)
        # set style
        self.m_stra_chart.setTitle("温差——时间折线图")  # 设置图题
        self.m_stra_chart.setTitleFont(QFont('SimHei', 15))  # 设置图题字体的类型和大小
        self.m_stra_chart.createDefaultAxes()  # 创建默认轴
        self.m_stra_chart.axisX().setTitleText('运行时间/h')  # 设置横坐标标题
        self.m_stra_chart.axisX().setLabelsFont(QFont('Times New Roman', 12))  # 设置横坐标刻度的字体类型和大小
        self.m_stra_chart.axisX().setTitleFont(QFont('SansSerif', 10))  # 设置横坐标标题字体的类型和大小
        self.m_stra_chart.axisY().setTitleText('温度/℃')  # 设置纵坐标标题
        self.m_stra_chart.axisY().setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        self.m_stra_chart.axisY().setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_stra_chart.legend().hide()
        self.ui.straView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.straView.setChart(self.m_stra_chart)
        self.ui.straView.show()

    def initBarPlot(self):
        '''
        初始化柱状图
        '''
        # ---整机部分---
        # 故障模式统计图
        self.m_patt_chart = QChart()
        self.m_patt_series = QBarSeries()
        axisY = QValueAxis()
        axisY.setTitleText('概率')  # 设置纵坐标标题
        axisY.setRange(0, 10e-3)  # 设置纵坐标范围
        axisY.setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        axisY.setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_patt_chart.setTitle("故障模式统计图")  # 设置图题
        self.m_patt_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_patt_chart.createDefaultAxes()  # 创建默认轴
        self.m_patt_chart.setAxisY(axisY)
        self.ui.pattView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.pattView.setChart(self.m_patt_chart)
        self.ui.pattView.show()
        # 故障部位统计图
        self.m_posi_chart = QChart()
        self.m_posi_series = QBarSeries()
        axisY = QValueAxis()
        axisY.setTitleText('概率')  # 设置纵坐标标题
        axisY.setRange(0, 10e-3)  # 设置纵坐标范围
        axisY.setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        axisY.setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_posi_chart.setTitle("故障部位统计图")  # 设置图题
        self.m_posi_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_posi_chart.createDefaultAxes()  # 创建默认轴
        self.m_posi_chart.setAxisY(axisY)
        self.ui.posiView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.posiView.setChart(self.m_posi_chart)
        self.ui.posiView.show()
        # 故障原因统计图
        self.m_reason_chart = QChart()
        self.m_reason_series = QBarSeries()
        axisY = QValueAxis()
        axisY.setTitleText('概率')  # 设置纵坐标标题
        axisY.setRange(0, 10e-3)  # 设置纵坐标范围
        axisY.setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        axisY.setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_reason_chart.setTitle("故障原因统计图")  # 设置图题
        self.m_reason_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_reason_chart.createDefaultAxes()  # 创建默认轴
        self.m_reason_chart.setAxisY(axisY)
        self.ui.reasonView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.reasonView.setChart(self.m_reason_chart)
        self.ui.reasonView.show()
        # 故障溯源统计图
        self.m_root_chart = QChart()
        self.m_root_series = QBarSeries()
        axisY = QValueAxis()
        axisY.setTitleText('概率')  # 设置纵坐标标题
        axisY.setRange(0, 10e-3)  # 设置纵坐标范围
        axisY.setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        axisY.setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_root_chart.setTitle("故障溯源统计图")  # 设置图题
        self.m_root_chart.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_root_chart.createDefaultAxes()  # 创建默认轴
        self.m_root_chart.setAxisY(axisY)
        self.ui.rootView.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.rootView.setChart(self.m_root_chart)
        self.ui.rootView.show()
        # ---子系统部分---
        self.m_patt_chart_2 = QChart()
        self.m_patt_series_2 = QBarSeries()
        axisY = QValueAxis()
        axisY.setTitleText('概率')  # 设置纵坐标标题
        axisY.setRange(0, 10e-3)  # 设置纵坐标范围
        axisY.setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        axisY.setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_patt_chart_2.setTitle("故障模式统计图")  # 设置图题
        self.m_patt_chart_2.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_patt_chart_2.createDefaultAxes()  # 创建默认轴
        self.m_patt_chart_2.setAxisY(axisY)
        self.ui.pattView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.pattView_2.setChart(self.m_patt_chart_2)
        self.ui.pattView_2.show()
        # 故障部位统计图
        self.m_posi_chart_2 = QChart()
        self.m_posi_series_2 = QBarSeries()
        axisY = QValueAxis()
        axisY.setTitleText('概率')  # 设置纵坐标标题
        axisY.setRange(0, 10e-3)  # 设置纵坐标范围
        axisY.setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        axisY.setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_posi_chart_2.setTitle("故障部位统计图")  # 设置图题
        self.m_posi_chart_2.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_posi_chart_2.createDefaultAxes()  # 创建默认轴
        self.m_posi_chart_2.setAxisY(axisY)
        self.ui.posiView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.posiView_2.setChart(self.m_posi_chart_2)
        self.ui.posiView_2.show()
        # 故障原因统计图
        self.m_reason_chart_2 = QChart()
        self.m_reason_series_2 = QBarSeries()
        axisY = QValueAxis()
        axisY.setTitleText('概率')  # 设置纵坐标标题
        axisY.setRange(0, 10e-3)  # 设置纵坐标范围
        axisY.setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        axisY.setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_reason_chart_2.setTitle("故障原因统计图")  # 设置图题
        self.m_reason_chart_2.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_reason_chart_2.createDefaultAxes()  # 创建默认轴
        self.m_reason_chart_2.setAxisY(axisY)
        self.ui.reasonView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.reasonView_2.setChart(self.m_reason_chart_2)
        self.ui.reasonView_2.show()
        # 故障溯源统计图
        self.m_root_chart_2 = QChart()
        self.m_root_series_2 = QBarSeries()
        axisY = QValueAxis()
        axisY.setTitleText('概率')  # 设置纵坐标标题
        axisY.setRange(0, 10e-3)  # 设置纵坐标范围
        axisY.setLabelsFont(QFont('Times New Roman', 12))  # 设置纵坐标刻度的字体类型和大小
        axisY.setTitleFont(QFont('SansSerif', 10))  # 设置纵坐标标题字体的类型和大小
        self.m_root_chart_2.setTitle("故障溯源统计图")  # 设置图题
        self.m_root_chart_2.setTitleFont(QFont('SimHei', 16))  # 设置图题字体的类型和大小
        self.m_root_chart_2.createDefaultAxes()  # 创建默认轴
        self.m_root_chart_2.setAxisY(axisY)
        self.ui.rootView_2.setRenderHint(QPainter.Antialiasing)  # 抗锯齿
        self.ui.rootView_2.setChart(self.m_root_chart_2)
        self.ui.rootView_2.show()

    def saveReport(self, save_dir):
        # 先保存所有结果到特定数据结构
        report = Report()
        item_head = ['尺度参数λ', '形状参数β', '可靠度', '失效率', 'MTBF']
        item_head_mt = ['尺度参数λ', '形状参数β', '修复率', '失效率', 'MTTR']
        item_head_ft = ['故障模式', '故障部位', '故障原因', '故障溯源']
        # 写入结果
        # 获取设备信息
        dev_id_1 = str(self.ui.devWidget.item(2).text()).strip()
        dev_name_1 = str(self.ui.devWidget.item(6).text()).strip()
        dev_num_1 = str(self.ui.devWidget.item(4).text()).strip()
        dev_id_2 = str(self.ui.devWidget_2.item(2).text()).strip()
        dev_name_2 = str(self.ui.devWidget_2.item(6).text()).strip()
        dev_num_2 = str(self.ui.devWidget_2.item(4).text()).strip()
        # 图片路径
        relia_env_temp_path = os.path.join(save_dir, '可靠度分析-环境温度折线图.jpg')
        relia_kni_temp_path = os.path.join(save_dir, '可靠度分析-刀头温度折线图.jpg')
        relia_rpa_path = os.path.join(save_dir, '可靠度分析-重复定位精度折线图.jpg')
        relia_temp_gap_path = os.path.join(save_dir, '可靠度分析-温差折线图.jpg')
        relia_scatter_path = os.path.join(save_dir, '可靠度分析-故障时间间隔散点图.jpg')
        relia_pdf_path = os.path.join(save_dir, '可靠度分析-威布尔分布概率密度图.jpg')
        relia_cdf_path = os.path.join(save_dir, '可靠度分析-威布尔分布累积分布图.jpg')
        relia_relia_path = os.path.join(save_dir, '可靠度分析-可靠度曲线图.jpg')
        relia_fail_path = os.path.join(save_dir, '可靠度分析-失效率曲线图.jpg')
        mt_scatter_path = os.path.join(save_dir, '维修性分析-维修时间散点图.jpg')
        mt_pdf_path = os.path.join(save_dir, '维修性分析-威布尔分布概率密度图.jpg')
        mt_cdf_path = os.path.join(save_dir, '维修性分析-威布尔分布累积分布图.jpg')
        mt_relia_path = os.path.join(save_dir, '维修性分析-修复率曲线图.jpg')
        mt_fail_path = os.path.join(save_dir, '维修性分析-失效率曲线图.jpg')
        fault_whole_patt_path = os.path.join(save_dir, '故障分析-整机-故障模式统计.jpg')
        fault_whole_posi_path = os.path.join(save_dir, '故障分析-整机-故障部位统计.jpg')
        fault_whole_reason_path = os.path.join(save_dir, '故障分析-整机-故障原因统计.jpg')
        fault_whole_root_path = os.path.join(save_dir, '故障分析-整机-故障溯源统计.jpg')
        fault_subsys_patt_path = os.path.join(save_dir, '故障分析-子系统-故障模式统计.jpg')
        fault_subsys_posi_path = os.path.join(save_dir, '故障分析-子系统-故障部位统计.jpg')
        fault_subsys_reason_path = os.path.join(save_dir, '故障分析-子系统-故障原因统计.jpg')
        fault_subsys_root_path = os.path.join(save_dir, '故障分析-子系统-故障溯源统计.jpg')
        # 保存表格
        if dev_id_1 != '' and dev_id_2 == '':  # 只有设备1
            report.setDevice(dev_id=int(dev_id_1), dev_name=dev_name_1, dev_num=dev_num_1)
            report.setReliaTable(key='设备名称', value=[dev_name_1])
            if self.ui.dev1ResultWidget.item(
                    Result.lamda_value.value).text().strip() != '' and self.ui.dev1MtResultWidget.item(
                Result.lamda_value.value).text().strip() != '' and self.ui.resultWidget_3.item(1).text().strip() != '':
                relia_lamda = float(self.ui.dev1ResultWidget.item(Result.lamda_value.value).text())
                relia_beta = float(self.ui.dev1ResultWidget.item(Result.beta_value.value).text())
                relia_result = [round(relia_lamda, 4), round(relia_beta, 4),
                                round(rf.reliability(relia_lamda, relia_beta, t=np.max(self.run_time_data_1)), 4),
                                round(rf.failure_rate(relia_lamda, relia_beta, t=np.max(self.run_time_data_1)), 4),
                                round(rf.mtbf(relia_lamda, relia_beta), 4)]
                for head, item in zip(item_head, relia_result):
                    report.setReliaTable(key=head, value=[item])
                report.setMaintainTable(key='设备名称', value=[dev_name_1])
                mt_lamda = float(self.ui.dev1MtResultWidget.item(Result.lamda_value.value).text())
                mt_beta = float(self.ui.dev1MtResultWidget.item(Result.beta_value.value).text())
                mt_result = [round(mt_lamda, 4), round(mt_beta, 4),
                             round(rf.reliability(mt_lamda, mt_beta, t=np.max(self.maintain_time_of_dev_1)), 4),
                             round(rf.failure_rate(mt_lamda, mt_beta, t=np.max(self.maintain_time_of_dev_1)), 4),
                             round(rf.mtbf(mt_lamda, mt_beta), 4)]
                for head, item in zip(item_head_mt, mt_result):
                    report.setMaintainTable(key=head, value=[item])
                report.setFaultTable(key='设备名称', value=[dev_name_1])
                for i in range(0, len(item_head_ft)):
                    report.setFaultTable(key=item_head_ft[i],
                                         value=[
                                             [re.findall(": (.+?) ", self.ui.resultWidget_3.item(2 * i + 1).text())[0],
                                              re.findall(": (.+?) ", self.ui.resultWidget_4.item(2 * i + 1).text())[
                                                  0]]])
            else:
                return Report()

        elif dev_id_1 == '' and dev_id_2 != '':  # 只有设备2
            report.setDevice(dev_id=int(dev_id_2), dev_name=dev_name_2, dev_num=dev_num_2)
            report.setReliaTable(key='设备名称', value=[dev_name_2])
            if self.ui.dev2ResultWidget.item(
                    Result.lamda_value.value).text().strip() != '' and self.ui.dev2MtResultWidget.item(
                Result.lamda_value.value).text().strip() != '' and self.ui.resultWidget_3.item(1).text().strip() != '':
                relia_lamda = float(self.ui.dev2ResultWidget.item(Result.lamda_value.value).text())
                relia_beta = float(self.ui.dev2ResultWidget.item(Result.beta_value.value).text())
                relia_result = [round(relia_lamda, 4), round(relia_beta, 4),
                                round(rf.reliability(relia_lamda, relia_beta, t=np.max(self.run_time_data_1)), 4),
                                round(rf.failure_rate(relia_lamda, relia_beta, t=np.max(self.run_time_data_2)), 4),
                                round(rf.mtbf(relia_lamda, relia_beta), 4)]
                for head, item in zip(item_head, relia_result):
                    report.setReliaTable(key=head, value=[item])
                report.setMaintainTable(key='设备名称', value=[dev_name_2])
                mt_lamda = float(self.ui.dev2MtResultWidget.item(Result.lamda_value.value).text())
                mt_beta = float(self.ui.dev2MtResultWidget.item(Result.beta_value.value).text())
                mt_result = [round(mt_lamda, 4), round(mt_beta, 4),
                             round(rf.reliability(mt_lamda, mt_beta, t=np.max(self.maintain_time_of_dev_2)), 4),
                             round(rf.failure_rate(mt_lamda, mt_beta, t=np.max(self.maintain_time_of_dev_2)), 4),
                             round(rf.mtbf(mt_lamda, mt_beta), 4)]
                for head, item in zip(item_head_mt, mt_result):
                    report.setMaintainTable(key=head, value=[item])
                report.setFaultTable(key='设备名称', value=[dev_name_2])
                for i in range(0, len(item_head_ft)):
                    report.setFaultTable(key=item_head_ft[i],
                                         value=[
                                             [re.findall(": (.+?) ", self.ui.resultWidget_3.item(2 * i + 1).text())[-1],
                                              re.findall(": (.+?) ", self.ui.resultWidget_4.item(2 * i + 1).text())[
                                                  -1]]])
            else:
                return Report()
        elif dev_id_1 != '' and dev_id_2 != '':
            report.setDevice(dev_id=int(dev_id_1), dev_name=dev_name_1, dev_num=dev_num_1)
            report.setDevice(dev_id=int(dev_id_2), dev_name=dev_name_2, dev_num=dev_num_2)
            report.setReliaTable(key='设备名称', value=[dev_name_1, dev_name_2])
            if self.ui.dev2ResultWidget.item(
                    Result.lamda_value.value).text().strip() != '' and self.ui.dev2MtResultWidget.item(
                Result.lamda_value.value).text().strip() != '' and self.ui.resultWidget_3.item(1).text().strip() != '':
                relia_lamda_dev_1 = float(self.ui.dev1ResultWidget.item(Result.lamda_value.value).text())
                relia_beta_dev_1 = float(self.ui.dev1ResultWidget.item(Result.beta_value.value).text())
                relia_result_dev_1 = [round(relia_lamda_dev_1, 4), round(relia_beta_dev_1, 4),
                                      round(rf.reliability(relia_lamda_dev_1, relia_beta_dev_1,
                                                           t=min(np.max(self.run_time_data_1),
                                                                 np.max(self.run_time_data_2))), 4),
                                      round(rf.failure_rate(relia_lamda_dev_1, relia_beta_dev_1,
                                                            t=min(np.max(self.run_time_data_1),
                                                                  np.max(self.run_time_data_2))), 4),
                                      round(rf.mtbf(relia_lamda_dev_1, relia_beta_dev_1), 4)]
                relia_lamda_dev_2 = float(self.ui.dev2ResultWidget.item(Result.lamda_value.value).text())
                relia_beta_dev_2 = float(self.ui.dev2ResultWidget.item(Result.beta_value.value).text())
                relia_result_dev_2 = [round(relia_lamda_dev_2, 4), round(relia_beta_dev_2, 4),
                                      round(rf.reliability(relia_lamda_dev_2, relia_beta_dev_2,
                                                           t=min(np.max(self.run_time_data_1),
                                                                 np.max(self.run_time_data_2))), 4),
                                      round(rf.failure_rate(relia_lamda_dev_2, relia_beta_dev_2,
                                                            t=min(np.max(self.run_time_data_1),
                                                                  np.max(self.run_time_data_2))), 4),
                                      round(rf.mtbf(relia_lamda_dev_2, relia_beta_dev_2), 4)]
                for i in range(len(item_head)):
                    report.setReliaTable(key=item_head[i], value=[relia_result_dev_1[i], relia_result_dev_2[i]])
                report.setMaintainTable(key='设备名称', value=[dev_name_1, dev_name_2])
                mt_lamda_dev_1 = float(self.ui.dev1MtResultWidget.item(Result.lamda_value.value).text())
                mt_beta_dev_1 = float(self.ui.dev1MtResultWidget.item(Result.beta_value.value).text())
                mt_result_dev_1 = [round(mt_lamda_dev_1, 4), round(mt_beta_dev_1, 4),
                                   round(rf.reliability(mt_lamda_dev_1, mt_beta_dev_1,
                                                        t=min(np.max(self.maintain_time_of_dev_1),
                                                              np.max(self.maintain_time_of_dev_2))), 4),
                                   round(rf.failure_rate(mt_lamda_dev_1, mt_beta_dev_1,
                                                         t=min(np.max(self.maintain_time_of_dev_1),
                                                               np.max(self.maintain_time_of_dev_2))), 4),
                                   round(rf.mtbf(mt_lamda_dev_1, mt_beta_dev_1), 4)]
                mt_lamda_dev_2 = float(self.ui.dev2MtResultWidget.item(Result.lamda_value.value).text())
                mt_beta_dev_2 = float(self.ui.dev2MtResultWidget.item(Result.beta_value.value).text())
                mt_result_dev_2 = [round(mt_lamda_dev_2, 4), round(mt_beta_dev_2, 4),
                                   round(rf.reliability(mt_lamda_dev_2, mt_beta_dev_2,
                                                        t=min(np.max(self.maintain_time_of_dev_1),
                                                              np.max(self.maintain_time_of_dev_2))), 4),
                                   round(rf.failure_rate(mt_lamda_dev_2, mt_beta_dev_2,
                                                         t=min(np.max(self.maintain_time_of_dev_1),
                                                               np.max(self.maintain_time_of_dev_2))), 4),
                                   round(rf.mtbf(mt_lamda_dev_2, mt_beta_dev_2), 4)]
                for i in range(len(item_head_mt)):
                    report.setMaintainTable(key=item_head_mt[i], value=[mt_result_dev_1[i], mt_result_dev_2[i]])
                report.setFaultTable(key='设备名称', value=[dev_name_1, dev_name_2])
                for i in range(0, len(item_head_ft)):
                    report.setFaultTable(key=item_head_ft[i],
                                         value=[
                                             [re.findall(": (.+?) ", self.ui.resultWidget_3.item(2 * i + 1).text())[0],
                                              re.findall(": (.+?) ", self.ui.resultWidget_4.item(2 * i + 1).text())[0]],
                                             [re.findall(": (.+?) ", self.ui.resultWidget_3.item(2 * i + 1).text())[-1],
                                              re.findall(": (.+?) ", self.ui.resultWidget_4.item(2 * i + 1).text())[-1]]
                                         ])
            else:
                return Report()
        else:
            return Report()
        # 保存图片路径
        report.setReliaImages('可靠度分析-环境温度折线图', relia_env_temp_path)
        report.setReliaImages('可靠度分析-刀头温度折线图', relia_kni_temp_path)
        report.setReliaImages('可靠度分析-重复定位精度折线图', relia_rpa_path)
        report.setReliaImages('可靠度分析-温差折线图', relia_temp_gap_path)
        report.setReliaImages('可靠度分析-故障时间间隔散点图', relia_scatter_path)
        report.setReliaImages('可靠度分析-威布尔分布概率密度图', relia_pdf_path)
        report.setReliaImages('可靠度分析-威布尔分布累积分布图', relia_cdf_path)
        report.setReliaImages('可靠度分析-可靠度曲线图', relia_relia_path)
        report.setReliaImages('可靠度分析-失效率曲线图', relia_fail_path)
        report.setMaintainImages('维修性分析-维修时间散点图', mt_scatter_path)
        report.setMaintainImages('维修性分析-威布尔分布概率密度图', mt_pdf_path)
        report.setMaintainImages('维修性分析-威布尔分布累积分布图', mt_cdf_path)
        report.setMaintainImages('维修性分析-修复率曲线图', mt_relia_path)
        report.setMaintainImages('维修性分析-失效率曲线图', mt_fail_path)
        report.setFaultImages('故障分析-整机-故障模式统计', fault_whole_patt_path)
        report.setFaultImages('故障分析-整机-故障部位统计', fault_whole_posi_path)
        report.setFaultImages('故障分析-整机-故障原因统计', fault_whole_reason_path)
        report.setFaultImages('故障分析-整机-故障溯源统计', fault_whole_root_path)
        report.setFaultImages('故障分析-子系统-故障模式统计', fault_subsys_patt_path)
        report.setFaultImages('故障分析-子系统-故障部位统计', fault_subsys_posi_path)
        report.setFaultImages('故障分析-子系统-故障原因统计', fault_subsys_reason_path)
        report.setFaultImages('故障分析-子系统-故障溯源统计', fault_subsys_root_path)
        delay_time = 100
        self.ui.stackedWidget.setCurrentIndex(0)
        self.ui.reliaWidget.setCurrentIndex(0)
        delay(delay_time)
        self.save_view(self.ui.tempView_1, relia_env_temp_path)
        self.save_view(self.ui.tempView_2, relia_kni_temp_path)
        self.save_view(self.ui.rpaView, relia_rpa_path)
        self.save_view(self.ui.straView, relia_temp_gap_path)
        self.ui.reliaWidget.setCurrentIndex(1)
        delay(delay_time)
        self.save_view(self.ui.scatterView, relia_scatter_path)
        self.ui.reliaWidget.setCurrentIndex(2)
        delay(delay_time)
        self.save_view(self.ui.pdfView, relia_pdf_path)
        self.save_view(self.ui.cdfView, relia_cdf_path)
        self.save_view(self.ui.reliaView, relia_relia_path)
        self.save_view(self.ui.failView, relia_fail_path)
        self.ui.stackedWidget.setCurrentIndex(1)
        self.ui.maintainWidget.setCurrentIndex(0)
        delay(delay_time)
        self.save_view(self.ui.scatterView_2, mt_scatter_path)
        self.ui.maintainWidget.setCurrentIndex(1)
        delay(delay_time)
        self.save_view(self.ui.pdfView_2, mt_pdf_path)
        self.save_view(self.ui.cdfView_2, mt_cdf_path)
        self.save_view(self.ui.reliaView_2, mt_relia_path)
        self.save_view(self.ui.failView_2, mt_fail_path)
        self.ui.stackedWidget.setCurrentIndex(2)
        self.ui.faultWidget.setCurrentIndex(0)
        delay(delay_time)
        self.save_view(self.ui.pattView, fault_whole_patt_path)
        self.save_view(self.ui.posiView, fault_whole_posi_path)
        self.save_view(self.ui.reasonView, fault_whole_reason_path)
        self.save_view(self.ui.rootView, fault_whole_root_path)
        self.ui.faultWidget.setCurrentIndex(1)
        delay(delay_time)
        self.save_view(self.ui.pattView_2, fault_subsys_patt_path)
        self.save_view(self.ui.posiView_2, fault_subsys_posi_path)
        self.save_view(self.ui.reasonView_2, fault_subsys_reason_path)
        self.save_view(self.ui.rootView_2, fault_subsys_root_path)
        self.ui.stackedWidget.setCurrentIndex(0)
        self.ui.reliaWidget.setCurrentIndex(0)
        return report

    def get_test_date(self):
        begin_date = '2000-01-01'
        end_date = '2000-01-01'
        query_sql = 'SELECT MIN(test_date), MAX(test_date) FROM record'
        self.query.prepare(query_sql)
        if not self.query.exec_():
            print(self.query.lastError().text())
        else:
            while self.query.next():
                begin_date = self.query.value(0)
                end_date = self.query.value(1)
        return [begin_date, end_date]

    def outputReport(self):
        dialog = OutputReportDialog()
        dialog.show()
        if dialog.exec_():
            person = dialog.get_person()
            save_dir = dialog.get_path()
            report = self.saveReport(save_dir)
            if not report.isEmpty():
                filename = '可靠性分析报告.docx'
                docx_file = os.path.join(save_dir, filename)
                document = Document()
                document.styles['Normal'].font.name = '宋体'
                document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                p = document.add_paragraph()  # 添加单位
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中
                run = p.add_run('单位：佛山市质量计量监督检测中心\n')
                run.font.bold = True
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(21)  # 字体大小设置，和word里面的字号相对应
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中设置
                # 添加标题
                run = p.add_run('可靠性测试分析报告\n')
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(21)  # 字体大小设置，和word里面的字号相对应
                # 添加测试信息
                run = p.add_run('\n测试人员：{0}\t\t\t测试日期：{1}\n'.format(person, '——'.join(self.get_test_date())))
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                # 添加报告内容
                ## 1.设备信息表格
                # 添加标题
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('1 设备信息')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(14)  # 字体大小设置，和word里面的字号相对应
                devices = report.getDevice()
                rows = len(devices)
                cols = 3
                table = document.add_table(rows=rows + 1, cols=cols, style='Table Grid')
                for i in range(rows):
                    tr = table.rows[i]._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "450")
                    trPr.append(trHeight)  # 表格高度设置
                col = table.columns[1]
                col.width = Inches(5)
                # 添加表头
                head = ['设备编号', '设备名称', '设备型号']
                heading_cells = table.rows[0].cells
                for i in range(cols):
                    p = heading_cells[i].paragraphs[0]
                    run = p.add_run(head[i])
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                    run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                    run.font.bold = True
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 填充表格
                for i in range(rows):
                    row_cells = table.rows[i + 1].cells
                    for j in range(cols):
                        p = row_cells[j].paragraphs[0]
                        run = p.add_run(str(devices[i][j]))
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                        run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                        run.font.bold = False
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ## 2.可靠性分析
                # 添加标题
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('\n2 可靠度分析')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(14)  # 字体大小设置，和word里面的字号相对应
                # 添加图
                p = document.add_paragraph()  # 添加标题
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('2.1 相关曲线图')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                relia_images = report.getReliaImages()
                document.add_picture(relia_images['可靠度分析-环境温度折线图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(relia_images['可靠度分析-刀头温度折线图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(relia_images['可靠度分析-重复定位精度折线图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(relia_images['可靠度分析-温差折线图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(relia_images['可靠度分析-故障时间间隔散点图'], width=Inches(5.60))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(relia_images['可靠度分析-威布尔分布概率密度图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(relia_images['可靠度分析-威布尔分布累积分布图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(relia_images['可靠度分析-可靠度曲线图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(relia_images['可靠度分析-可靠度曲线图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                # 添加表格
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('2.2 分析结果表格')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                relia_table = report.getReliaTable()
                # 添加表头
                head = ['设备名称', '尺度参数λ', '形状参数β', '可靠度', '失效率', 'MTBF']
                rows = len(relia_table[head[0]])
                cols = len(head)
                table = document.add_table(rows=rows + 1, cols=cols, style='Table Grid')
                for i in range(rows):
                    tr = table.rows[i]._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "450")
                    trPr.append(trHeight)  # 表格高度设置
                col = table.columns[1]
                col.width = Inches(5)
                heading_cells = table.rows[0].cells
                for i in range(cols):
                    p = heading_cells[i].paragraphs[0]
                    run = p.add_run(head[i])
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                    run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                    run.font.bold = True
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 填充表格
                for i in range(rows):
                    row_cells = table.rows[i + 1].cells
                    for j in range(cols):
                        p = row_cells[j].paragraphs[0]
                        run = p.add_run(str(relia_table[head[j]][i]))
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                        run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                        run.font.bold = False
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ## 3.维修性分析
                # 添加标题
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('\n3 维修性分析')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(14)  # 字体大小设置，和word里面的字号相对应
                # 添加图
                p = document.add_paragraph()  # 添加标题
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('3.1 相关曲线图')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                mt_images = report.getMaintainImages()
                document.add_picture(mt_images['维修性分析-维修时间散点图'], width=Inches(5.60))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(mt_images['维修性分析-威布尔分布概率密度图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(mt_images['维修性分析-威布尔分布累积分布图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(mt_images['维修性分析-修复率曲线图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(mt_images['维修性分析-失效率曲线图'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                # 添加表格
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('3.2 分析结果表格')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                mt_table = report.getMaintainTable()
                # 添加表头
                head = ['设备名称', '尺度参数λ', '形状参数β', '修复率', '失效率', 'MTTR']
                rows = len(mt_table[head[0]])
                cols = len(head)
                table = document.add_table(rows=rows + 1, cols=cols, style='Table Grid')
                for i in range(rows):
                    tr = table.rows[i]._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "450")
                    trPr.append(trHeight)  # 表格高度设置
                col = table.columns[1]
                col.width = Inches(5)
                heading_cells = table.rows[0].cells
                for i in range(cols):
                    p = heading_cells[i].paragraphs[0]
                    run = p.add_run(head[i])
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                    run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                    run.font.bold = True
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 填充表格
                for i in range(rows):
                    row_cells = table.rows[i + 1].cells
                    for j in range(cols):
                        p = row_cells[j].paragraphs[0]
                        run = p.add_run(str(mt_table[head[j]][i]))
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                        run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                        run.font.bold = False
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                ## 4.故障分析
                # 添加标题
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('\n4 故障分析')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(14)  # 字体大小设置，和word里面的字号相对应
                # 整机部分
                p = document.add_paragraph()  # 添加标题
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('4.1 整机分析')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                # 添加图
                p = document.add_paragraph()  # 添加标题
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('4.1.1 相关统计图')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                ft_images = report.getFaultImages()
                document.add_picture(ft_images['故障分析-整机-故障模式统计'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(ft_images['故障分析-整机-故障部位统计'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(ft_images['故障分析-整机-故障原因统计'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(ft_images['故障分析-整机-故障溯源统计'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                # 添加表格
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('4.1.2 分析结果表格')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                ft_table = report.getFaultTable()
                # 添加表头
                head = ['设备名称', '故障模式', '故障部位', '故障原因', '故障溯源']
                rows = len(ft_table[head[0]])
                cols = len(head)
                table = document.add_table(rows=rows + 1, cols=cols, style='Table Grid')
                for i in range(rows):
                    tr = table.rows[i]._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "450")
                    trPr.append(trHeight)  # 表格高度设置
                col = table.columns[1]
                col.width = Inches(5)
                heading_cells = table.rows[0].cells
                for i in range(cols):
                    p = heading_cells[i].paragraphs[0]
                    run = p.add_run(head[i])
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                    run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                    run.font.bold = True
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 填充表格
                for i in range(rows):
                    row_cells = table.rows[i + 1].cells
                    for j in range(cols):
                        p = row_cells[j].paragraphs[0]
                        if j == 0:
                            run = p.add_run(str(ft_table[head[j]][i]))
                        else:
                            run = p.add_run(str(ft_table[head[j]][i][0]))
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                        run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                        run.font.bold = False
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 子系统部分
                p = document.add_paragraph()  # 添加标题
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('\n4.2 整机分析')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                # 添加图
                p = document.add_paragraph()  # 添加标题
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('4.2.1 相关统计图')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                document.add_picture(ft_images['故障分析-子系统-故障模式统计'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(ft_images['故障分析-子系统-故障部位统计'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(ft_images['故障分析-子系统-故障原因统计'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                document.add_picture(ft_images['故障分析-子系统-故障溯源统计'])
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 图片居中设置
                # 添加表格
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题靠左
                run = p.add_run('4.2.2 分析结果表格')
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                ft_table = report.getFaultTable()
                # 添加表头
                table = document.add_table(rows=rows + 1, cols=cols, style='Table Grid')
                for i in range(rows):
                    tr = table.rows[i]._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), "450")
                    trPr.append(trHeight)  # 表格高度设置
                col = table.columns[1]
                col.width = Inches(5)
                heading_cells = table.rows[0].cells
                for i in range(cols):
                    p = heading_cells[i].paragraphs[0]
                    run = p.add_run(head[i])
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                    run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                    run.font.bold = True
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 填充表格
                for i in range(rows):
                    row_cells = table.rows[i + 1].cells
                    for j in range(cols):
                        p = row_cells[j].paragraphs[0]
                        if j == 0:
                            run = p.add_run(str(ft_table[head[j]][i]))
                        else:
                            run = p.add_run(str(ft_table[head[j]][i][1]))
                        run.font.color.rgb = RGBColor(0, 0, 0)  # 颜色设置，这里是用RGB颜色
                        run.font.size = Pt(12)  # 字体大小设置，和word里面的字号相对应
                        run.font.bold = False
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 保存文档
                document.save(docx_file)
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '保存成功',
                                      QMessageBox.Ok)
                msg_box.exec_()
                for filename in os.listdir(save_dir):
                    if '.jpg' in filename:
                        os.remove(os.path.join(save_dir, filename))
            else:
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '无分析结果',
                                      QMessageBox.Ok)
                msg_box.exec_()

    def on_listWidget_itemClicked(self):
        curr_row = self.ui.listWidget.currentRow()
        self.ui.stackedWidget.setCurrentIndex(curr_row)

    def read_raw_data(self, dev_id):
        '''
        读取设备dev_id的运行记录
        :param dev_id:
        :return:
        '''
        data_table = dict()  # 保存读取数据
        # 数据库查询
        x_table = []
        run_times = []
        env_temps = []
        kni_temps = []
        rpas = []
        query_sql = 'select run_time, env_temp, kni_temp, rpa from record where dev_id = :dev_id and axis = :axis'
        self.query.prepare(query_sql)
        self.query.bindValue(':dev_id', dev_id)
        self.query.bindValue(':axis', 'X')
        if not self.query.exec_():
            print(self.query.lastError().text())
            return data_table
        else:
            while self.query.next():
                run_times.append(self.query.value(RUN_TIME))
                env_temps.append(self.query.value(ENV_TEMP))
                kni_temps.append(self.query.value(KNI_TEMP))
                rpas.append(self.query.value(RPA))
            x_table.append(run_times)
            x_table.append(env_temps)
            x_table.append(kni_temps)
            x_table.append(rpas)
        y_table = []
        run_times = []
        env_temps = []
        kni_temps = []
        rpas = []
        query_sql = 'select run_time, env_temp, kni_temp, rpa from record where dev_id = :dev_id and axis = :axis'
        self.query.prepare(query_sql)
        self.query.bindValue(':dev_id', dev_id)
        self.query.bindValue(':axis', 'Y')
        if not self.query.exec_():
            print(self.query.lastError().text())
            return data_table
        else:
            while self.query.next():
                run_times.append(self.query.value(RUN_TIME))
                env_temps.append(self.query.value(ENV_TEMP))
                kni_temps.append(self.query.value(KNI_TEMP))
                rpas.append(self.query.value(RPA))
            y_table.append(run_times)
            y_table.append(env_temps)
            y_table.append(kni_temps)
            y_table.append(rpas)
        z_table = []
        run_times = []
        env_temps = []
        kni_temps = []
        rpas = []
        query_sql = 'select run_time, env_temp, kni_temp, rpa from record where dev_id = :dev_id and axis = :axis'
        self.query.prepare(query_sql)
        self.query.bindValue(':dev_id', dev_id)
        self.query.bindValue(':axis', 'Z')
        if not self.query.exec_():
            print(self.query.lastError().text())
            return data_table
        else:
            while self.query.next():
                run_times.append(self.query.value(RUN_TIME))
                env_temps.append(self.query.value(ENV_TEMP))
                kni_temps.append(self.query.value(KNI_TEMP))
                rpas.append(self.query.value(RPA))
            z_table.append(run_times)
            z_table.append(env_temps)
            z_table.append(kni_temps)
            z_table.append(rpas)
        # 存储
        data_table['x_table'] = x_table
        data_table['y_table'] = y_table
        data_table['z_table'] = z_table

        return data_table

    def find_fault_time(self, data_table, temp_thresh, rpa_thresh):
        '''
        统计故障时间间隔
        :param data_table: 数据表
        :param temp_thresh: 温差阈值
        :param rpa_thresh: 精度阈值
        :return:
        '''
        fault_time = []
        x_table = data_table['x_table']
        y_table = data_table['y_table']
        z_table = data_table['z_table']
        data_num = len(x_table[0])
        for i in range(0, data_num):
            if (x_table[KNI_TEMP][i] - x_table[ENV_TEMP][i] > temp_thresh or x_table[RPA][i] < rpa_thresh) or \
                    (y_table[KNI_TEMP][i] - y_table[ENV_TEMP][i] > temp_thresh or y_table[RPA][i] < rpa_thresh) or \
                    (z_table[KNI_TEMP][i] - z_table[ENV_TEMP][i] > temp_thresh or z_table[RPA][i] < rpa_thresh):
                fault_time.append(x_table[RUN_TIME][i])
        return fault_time

    def compute_curves(self, run_time_data, method=0):
        t_max = max(run_time_data)
        t_range = np.arange(1, t_max, 1)
        lambda_hat, beta_hat, beta_std = estimate.least_square(run_time_data)
        if method == 0:
            ## 以下是最小二乘法计算
            # 计算概率密度曲线
            pdf_prob_ls = rf.Weibull_pdf(lambda_hat, beta_hat, t_range)
            # 计算累积分布曲线
            cdf_prob_ls = rf.Weibull_cdf(lambda_hat, beta_hat, t_range)
            # 计算可靠度曲线
            relia_prob_ls = rf.reliability(lambda_hat, beta_hat, t_range)
            # 计算失效率曲线
            fail_prob_ls = rf.failure_rate(lambda_hat, beta_hat, t_range)
            return lambda_hat, beta_hat, pdf_prob_ls, cdf_prob_ls, relia_prob_ls, fail_prob_ls
        else:
            ## 以下是极大后验估计法计算
            lambda_map, beta_map = estimate.map_estimate([lambda_hat, beta_hat], t=run_time_data, mu=lambda_hat,
                                                         sigma=beta_std)
            # 计算概率密度曲线
            pdf_prob_map = rf.Weibull_pdf(lambda_map, beta_map, t_range)
            # 计算累积分布曲线
            cdf_prob_map = rf.Weibull_cdf(lambda_map, beta_map, t_range)
            # 计算可靠度曲线
            relia_prob_map = rf.reliability(lambda_map, beta_map, t_range)
            # 计算失效率曲线
            fail_prob_map = rf.failure_rate(lambda_map, beta_map, t_range)
            return lambda_map, beta_map, pdf_prob_map, cdf_prob_map, relia_prob_map, fail_prob_map

    def plot_relia_charts(self, run_time_data, m_scatter_chart, m_scatter_series, m_pdf_chart, m_pdf_series_ls,
                          m_pdf_series_map, m_cdf_chart, m_cdf_series_ls, m_cdf_series_map, m_relia_chart,
                          m_relia_series_ls, m_relia_series_map, m_fail_chart, m_fail_series_ls, m_fail_series_map,
                          resultWidget):
        '''
        绘制可靠性分析相关曲线
        :param run_time_data_1: 设备1的故障时间间隔
        :param run_time_data_2: 设备2的故障时间间隔
        :return:
        '''
        t_max = np.max(run_time_data)
        t_range = np.arange(1, t_max, 1)
        run_time_prob = np.array(
            [(i - 0.3) / (len(run_time_data) + 0.4) for i in range(1, len(run_time_data) + 1)])
        # 绘制散点图
        for t, prob in zip(run_time_data, run_time_prob):
            m_scatter_series.append(t, prob)
        # 重新设定x轴范围
        old_max = m_scatter_chart.axisX().max()
        m_scatter_chart.axisX().setRange(0, np.max([t_max, old_max]))
        m_pdf_chart.axisX().setRange(0, np.max([t_max, old_max]))
        m_cdf_chart.axisX().setRange(0, np.max([t_max, old_max]))
        m_relia_chart.axisX().setRange(0, np.max([t_max, old_max]))
        m_fail_chart.axisX().setRange(0, np.max([t_max, old_max]))
        # 计算各曲线
        lambda_hat, beta_hat, pdf_prob_ls, cdf_prob_ls, relia_prob_ls, fail_prob_ls = self.compute_curves(run_time_data,
                                                                                                          method=0)
        lambda_map, beta_map, pdf_prob_map, cdf_prob_map, relia_prob_map, fail_prob_map = self.compute_curves(
            run_time_data, method=1)
        # 保存全局变量
        if m_pdf_series_ls == self.m_pdf_series_ls:
            self.lambda_hat_1, self.beta_hat_1 = lambda_hat, beta_hat
        else:
            self.lambda_hat_2, self.beta_hat_2 = lambda_hat, beta_hat
        if m_pdf_series_map == self.m_pdf_series_map:
            self.lambda_map_1, self.beta_map_1 = lambda_map, beta_map
        else:
            self.lambda_map_2, self.beta_map_2 = lambda_map, beta_map
        # 重新设定y轴范围
        m_pdf_chart.axisY().setRange(0, max(np.max(pdf_prob_ls), np.max(pdf_prob_map), m_pdf_chart.axisY().max()))
        m_cdf_chart.axisY().setRange(0, max(np.max(cdf_prob_ls), np.max(cdf_prob_map), m_cdf_chart.axisY().max()))
        m_relia_chart.axisY().setRange(0,
                                       max(np.max(relia_prob_ls), np.max(relia_prob_map), m_relia_chart.axisY().max()))
        m_fail_chart.axisY().setRange(0, max(np.max(fail_prob_ls), np.max(fail_prob_map), m_fail_chart.axisY().max()))
        method = self.ui.comboBox.currentIndex()  # 获取算法选择索引, 0---最小二乘法, 1---贝叶斯估计法
        if method == 0:
            m_pdf_chart.removeSeries(self.m_pdf_series_map)
            m_cdf_chart.removeSeries(self.m_cdf_series_map)
            m_relia_chart.removeSeries(self.m_relia_series_map)
            m_fail_chart.removeSeries(self.m_fail_series_map)

            m_pdf_chart.removeSeries(self.m_pdf_series_map_2)
            m_cdf_chart.removeSeries(self.m_cdf_series_map_2)
            m_relia_chart.removeSeries(self.m_relia_series_map_2)
            m_fail_chart.removeSeries(self.m_fail_series_map_2)
            # 填充表格
            resultWidget.item(Result.lamda_value.value).setText('%.4f' % lambda_hat)
            resultWidget.item(Result.beta_value.value).setText('%.4f' % beta_hat)
            resultWidget.item(Result.mtbf_value.value).setText('%.4f' % rf.mtbf(lambda_hat, beta_hat))
        else:
            m_pdf_chart.removeSeries(self.m_pdf_series_ls)
            m_cdf_chart.removeSeries(self.m_cdf_series_ls)
            m_relia_chart.removeSeries(self.m_relia_series_ls)
            m_fail_chart.removeSeries(self.m_fail_series_ls)

            m_pdf_chart.removeSeries(self.m_pdf_series_ls_2)
            m_cdf_chart.removeSeries(self.m_cdf_series_ls_2)
            m_relia_chart.removeSeries(self.m_relia_series_ls_2)
            m_fail_chart.removeSeries(self.m_fail_series_ls_2)
            # 填充表格
            resultWidget.item(Result.lamda_value.value).setText('%.4f' % lambda_map)
            resultWidget.item(Result.beta_value.value).setText('%.4f' % beta_map)
            resultWidget.item(Result.mtbf_value.value).setText('%.4f' % rf.mtbf(lambda_map, beta_map))
        # 绘制概率密度曲线
        m_pdf_series_ls.clear()
        for t, prob in zip(t_range, pdf_prob_ls):
            m_pdf_series_ls.append(t, prob)
        # 绘制累积分布曲线
        m_cdf_series_ls.clear()
        for t, prob in zip(t_range, cdf_prob_ls):
            m_cdf_series_ls.append(t, prob)
        # 绘制可靠度曲线
        m_relia_series_ls.clear()
        for t, prob in zip(t_range, relia_prob_ls):
            m_relia_series_ls.append(t, prob)
        # 绘制失效率曲线
        m_fail_series_ls.clear()
        for t, prob in zip(t_range, fail_prob_ls):
            m_fail_series_ls.append(t, prob)
        # 绘制概率密度曲线
        m_pdf_series_map.clear()
        for t, prob in zip(t_range, pdf_prob_map):
            m_pdf_series_map.append(t, prob)
        # 绘制累积分布曲线
        m_cdf_series_map.clear()
        for t, prob in zip(t_range, cdf_prob_map):
            m_cdf_series_map.append(t, prob)
        # 绘制可靠度曲线
        m_relia_series_map.clear()
        for t, prob in zip(t_range, relia_prob_map):
            m_relia_series_map.append(t, prob)
        # 绘制失效率曲线
        m_fail_series_map.clear()
        for t, prob in zip(t_range, fail_prob_map):
            m_fail_series_map.append(t, prob)

    def plot_raw_data(self, data_table, m_env_temp_series, m_env_temp_scatter, m_kni_temp_series, m_kni_temp_scatter,
                      m_rpa_series, m_rpa_scatter, m_stra_series, m_stra_scatter):
        if len(data_table) > 0:
            record_num = len(data_table[0])
            # 绘制温度曲线
            t_max = max(data_table[RUN_TIME])
            env_temp_max = max(data_table[ENV_TEMP])
            env_temp_min = min(data_table[ENV_TEMP])
            self.m_env_temp_chart.axisX().setRange(0, max(t_max, self.m_env_temp_chart.axisX().max()))
            lower = env_temp_min if self.m_env_temp_chart.axisY().min() == 0 else min(env_temp_min,
                                                                                      self.m_env_temp_chart.axisY().min())
            self.m_env_temp_chart.axisY().setRange(lower, max(env_temp_max, self.m_env_temp_chart.axisY().max()))
            self.m_env_temp_series.clear()
            for i in range(0, record_num):
                t = data_table[RUN_TIME][i]
                temp1 = data_table[ENV_TEMP][i]
                m_env_temp_series.append(t, temp1)
                m_env_temp_scatter.append(t, temp1)
            # 绘制刀头温度曲线
            kni_temp_max = max(data_table[KNI_TEMP])
            kni_temp_min = min(data_table[KNI_TEMP])
            self.m_kni_temp_chart.axisX().setRange(0, max(t_max, self.m_kni_temp_chart.axisX().max()))
            lower = kni_temp_min if self.m_kni_temp_chart.axisY().min() == 0 else min(kni_temp_min,
                                                                                      self.m_kni_temp_chart.axisY().min())
            self.m_kni_temp_chart.axisY().setRange(lower, max(kni_temp_max, self.m_kni_temp_chart.axisY().max()))
            m_kni_temp_series.clear()
            for i in range(0, record_num):
                t = data_table[RUN_TIME][i]
                temp2 = data_table[KNI_TEMP][i]
                m_kni_temp_series.append(t, temp2)
                m_kni_temp_scatter.append(t, temp2)
            # 绘制重复定位精度曲线
            rpa_data_list = data_table[RPA]
            rpa_max = max(rpa_data_list)
            rpa_min = min(rpa_data_list)
            self.m_rpa_chart.axisX().setRange(0, max(t_max, self.m_rpa_chart.axisX().max()))
            lower = rpa_min if self.m_rpa_chart.axisY().min() == 0 else min(rpa_min, self.m_rpa_chart.axisY().min())
            self.m_rpa_chart.axisY().setRange(lower, max(rpa_max, self.m_rpa_chart.axisY().max()))
            m_rpa_series.clear()
            for i in range(0, record_num):
                t = data_table[RUN_TIME][i]
                rpa = rpa_data_list[i]
                m_rpa_series.append(t, rpa)
                m_rpa_scatter.append(t, rpa)
            # 绘制温差曲线
            temp_gap_list = np.array(data_table[KNI_TEMP]) - np.array(data_table[ENV_TEMP])
            temp_gap_max = max(temp_gap_list)
            temp_gap_min = min(temp_gap_list)
            self.m_stra_chart.axisX().setRange(0, max(t_max, self.m_stra_chart.axisX().max()))
            lower = temp_gap_min if self.m_stra_chart.axisY().min() == 0 else min(temp_gap_min,
                                                                                  self.m_stra_chart.axisY().min())
            self.m_stra_chart.axisY().setRange(lower, max(temp_gap_max, self.m_stra_chart.axisY().max()))
            m_stra_series.clear()
            for i in range(0, record_num):
                t = data_table[RUN_TIME][i]
                temp_gap = temp_gap_list[i]
                m_stra_series.append(t, temp_gap)
                m_stra_scatter.append(t, temp_gap)

    @pyqtSlot()
    def on_loadButton_clicked(self):
        from utils.valid_utils import isNumber

        def check_data_exist(data_table):
            if len(data_table) <= 0 or (
                    data_table['x_table'] == [[], [], [], []] and data_table['y_table'] == [[], [], [], []] and
                    data_table['z_table'] == [[], [], [], []]):
                flag = False
            else:
                flag = True
            return flag

        temp_thresh = self.ui.tempEdit.text()
        rpa_thresh = self.ui.precEdit.text()
        if (temp_thresh == '' or rpa_thresh == '') or (' ' in temp_thresh or ' ' in rpa_thresh):
            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '请填写故障判断标准',
                                  QMessageBox.Ok)
            msg_box.exec_()
            return
        if not (isNumber(temp_thresh) and isNumber(rpa_thresh)):
            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '请填写正确数值',
                                  QMessageBox.Ok)
            msg_box.exec_()
            return
        temp_thresh = float(temp_thresh)
        rpa_thresh = float(rpa_thresh)
        dev_1 = self.ui.devWidget.item(2).text()
        dev_2 = self.ui.devWidget_2.item(2).text()
        if dev_1.strip(' ') != '' and dev_2.strip(' ') == '':  # 只有设备1的数据
            dev_1 = int(dev_1)
            self.data_table_of_dev_1 = self.read_raw_data(dev_1)
            if not check_data_exist(self.data_table_of_dev_1):
                msg_box = QMessageBox(QMessageBox.Warning,
                                      '提示',
                                      '数据库无记录！',
                                      QMessageBox.Ok)
                msg_box.exec_()
                return
            else:
                # ---原始数据可视化---
                # first clear old series
                self.m_env_temp_series.clear()
                self.m_env_temp_scatter.clear()
                self.m_kni_temp_series.clear()
                self.m_kni_temp_scatter.clear()
                self.m_rpa_series.clear()
                self.m_rpa_scatter.clear()
                self.m_stra_series.clear()
                self.m_stra_scatter.clear()
                self.m_env_temp_series_2.clear()
                self.m_env_temp_scatter_2.clear()
                self.m_kni_temp_series_2.clear()
                self.m_kni_temp_scatter_2.clear()
                self.m_rpa_series_2.clear()
                self.m_rpa_scatter_2.clear()
                self.m_stra_series_2.clear()
                self.m_stra_scatter_2.clear()
                axis = self.ui.axisComboBox.currentIndex()  # 0---x axis, 1---y axis, 2---z axis
                if axis == 0:
                    self.plot_raw_data(self.data_table_of_dev_1['x_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                elif axis == 1:
                    self.plot_raw_data(self.data_table_of_dev_1['y_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                else:
                    self.plot_raw_data(self.data_table_of_dev_1['z_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                # ---可靠性分析相关---
                # 读取数据
                fault_time = self.find_fault_time(self.data_table_of_dev_1, temp_thresh, rpa_thresh)
                if len(fault_time) <= 0:  # 没有统计出故障，只需提示
                    msg_box = QMessageBox(QMessageBox.Information,
                                          '提示',
                                          '设备无故障',
                                          QMessageBox.Ok)
                    msg_box.exec_()
                else:
                    run_time_data_1 = np.array(fault_time)
                    self.run_time_data_1 = run_time_data_1
                    self.plot_relia_charts(run_time_data_1, self.m_scatter_chart, self.m_scatter_series,
                                           self.m_pdf_chart, self.m_pdf_series_ls, self.m_pdf_series_map,
                                           self.m_cdf_chart, self.m_cdf_series_ls, self.m_cdf_series_map,
                                           self.m_relia_chart, self.m_relia_series_ls, self.m_relia_series_map,
                                           self.m_fail_chart, self.m_fail_series_ls, self.m_fail_series_map,
                                           self.ui.dev1ResultWidget)
        elif dev_1.strip(' ') == '' and dev_2.strip(' ') != '':
            dev_2 = int(dev_2)
            self.data_table_of_dev_2 = self.read_raw_data(dev_2)
            if not check_data_exist(self.data_table_of_dev_2):
                msg_box = QMessageBox(QMessageBox.Warning,
                                      '提示',
                                      '数据库无记录！',
                                      QMessageBox.Ok)
                msg_box.exec_()
                return
            else:
                # ---原始数据可视化---
                # first clear old series
                self.m_env_temp_series.clear()
                self.m_env_temp_scatter.clear()
                self.m_kni_temp_series.clear()
                self.m_kni_temp_scatter.clear()
                self.m_rpa_series.clear()
                self.m_rpa_scatter.clear()
                self.m_stra_series.clear()
                self.m_stra_scatter.clear()
                self.m_env_temp_series_2.clear()
                self.m_env_temp_scatter_2.clear()
                self.m_kni_temp_series_2.clear()
                self.m_kni_temp_scatter_2.clear()
                self.m_rpa_series_2.clear()
                self.m_rpa_scatter_2.clear()
                self.m_stra_series_2.clear()
                self.m_stra_scatter_2.clear()
                axis = self.ui.axisComboBox.currentIndex()  # 0---x axis, 1---y axis, 2---z axis
                if axis == 0:
                    self.plot_raw_data(self.data_table_of_dev_2['x_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                elif axis == 1:
                    self.plot_raw_data(self.data_table_of_dev_2['y_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                else:
                    self.plot_raw_data(self.data_table_of_dev_2['z_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                # ---可靠性分析相关---
                # 读取数据
                fault_time = self.find_fault_time(self.data_table_of_dev_2, temp_thresh, rpa_thresh)
                if len(fault_time) <= 0:  # 没有统计出故障，只需提示
                    msg_box = QMessageBox(QMessageBox.Information,
                                          '提示',
                                          '设备无故障',
                                          QMessageBox.Ok)
                    msg_box.exec_()
                else:
                    run_time_data_2 = np.array(fault_time)
                    self.run_time_data_2 = run_time_data_2
                    self.plot_relia_charts(run_time_data_2, self.m_scatter_chart, self.m_scatter_series,
                                           self.m_pdf_chart, self.m_pdf_series_ls_2, self.m_pdf_series_map_2,
                                           self.m_cdf_chart, self.m_cdf_series_ls_2, self.m_cdf_series_map_2,
                                           self.m_relia_chart, self.m_relia_series_ls_2, self.m_relia_series_map_2,
                                           self.m_fail_chart, self.m_fail_series_ls_2, self.m_fail_series_map_2,
                                           self.ui.dev2ResultWidget)
        elif dev_1.strip(' ') != '' and dev_2.strip(' ') != '':
            dev_1 = int(dev_1)
            dev_2 = int(dev_2)
            self.data_table_of_dev_1 = self.read_raw_data(dev_1)
            self.data_table_of_dev_2 = self.read_raw_data(dev_2)
            if not (check_data_exist(self.data_table_of_dev_1) or check_data_exist(self.data_table_of_dev_2)):
                msg_box = QMessageBox(QMessageBox.Warning,
                                      '提示',
                                      '数据库无记录！',
                                      QMessageBox.Ok)
                msg_box.exec_()
                return
            else:
                # ---原始数据可视化---
                # first clear old series
                self.m_env_temp_series.clear()
                self.m_env_temp_scatter.clear()
                self.m_kni_temp_series.clear()
                self.m_kni_temp_scatter.clear()
                self.m_rpa_series.clear()
                self.m_rpa_scatter.clear()
                self.m_stra_series.clear()
                self.m_stra_scatter.clear()
                self.m_env_temp_series_2.clear()
                self.m_env_temp_scatter_2.clear()
                self.m_kni_temp_series_2.clear()
                self.m_kni_temp_scatter_2.clear()
                self.m_rpa_series_2.clear()
                self.m_rpa_scatter_2.clear()
                self.m_stra_series_2.clear()
                self.m_stra_scatter_2.clear()
                axis = self.ui.axisComboBox.currentIndex()  # 0---x axis, 1---y axis, 2---z axis
                if axis == 0:
                    self.plot_raw_data(self.data_table_of_dev_1['x_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                    self.plot_raw_data(self.data_table_of_dev_2['x_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                elif axis == 1:
                    self.plot_raw_data(self.data_table_of_dev_1['y_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                    self.plot_raw_data(self.data_table_of_dev_2['y_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                else:
                    self.plot_raw_data(self.data_table_of_dev_1['z_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                    self.plot_raw_data(self.data_table_of_dev_2['z_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                # ---可靠性分析相关---
                # 读取数据
                fault_time_1 = self.find_fault_time(self.data_table_of_dev_1, temp_thresh, rpa_thresh)
                fault_time_2 = self.find_fault_time(self.data_table_of_dev_2, temp_thresh, rpa_thresh)
                if len(fault_time_2) <= 0 and len(fault_time_1) > 0:  # 没有统计出故障，只需提示
                    msg_box = QMessageBox(QMessageBox.Information,
                                          '提示',
                                          '设备2无故障',
                                          QMessageBox.Ok)
                    msg_box.exec_()
                    run_time_data_1 = np.array(fault_time_1)
                    self.plot_relia_charts(run_time_data_1, self.m_scatter_chart, self.m_scatter_series,
                                           self.m_pdf_chart, self.m_pdf_series_ls, self.m_pdf_series_map,
                                           self.m_cdf_chart, self.m_cdf_series_ls, self.m_cdf_series_map,
                                           self.m_relia_chart, self.m_relia_series_ls, self.m_relia_series_map,
                                           self.m_fail_chart, self.m_fail_series_ls, self.m_fail_series_map,
                                           self.ui.dev1ResultWidget)
                elif len(fault_time_2) > 0 and len(fault_time_1) <= 0:
                    msg_box = QMessageBox(QMessageBox.Information,
                                          '提示',
                                          '设备1无故障',
                                          QMessageBox.Ok)
                    msg_box.exec_()
                    run_time_data_2 = np.array(fault_time_2)
                    self.plot_relia_charts(run_time_data_2, self.m_scatter_chart, self.m_scatter_series_2,
                                           self.m_pdf_chart, self.m_pdf_series_ls_2, self.m_pdf_series_map_2,
                                           self.m_cdf_chart, self.m_cdf_series_ls_2, self.m_cdf_series_map_2,
                                           self.m_relia_chart, self.m_relia_series_ls_2, self.m_relia_series_map_2,
                                           self.m_fail_chart, self.m_fail_series_ls_2, self.m_fail_series_map_2,
                                           self.ui.dev2ResultWidget)
                elif len(fault_time_2) > 0 and len(fault_time_1) > 0:
                    run_time_data_1 = np.array(fault_time_1)
                    self.plot_relia_charts(run_time_data_1, self.m_scatter_chart, self.m_scatter_series,
                                           self.m_pdf_chart, self.m_pdf_series_ls, self.m_pdf_series_map,
                                           self.m_cdf_chart, self.m_cdf_series_ls, self.m_cdf_series_map,
                                           self.m_relia_chart, self.m_relia_series_ls, self.m_relia_series_map,
                                           self.m_fail_chart, self.m_fail_series_ls, self.m_fail_series_map,
                                           self.ui.dev1ResultWidget)
                    run_time_data_2 = np.array(fault_time_2)
                    self.plot_relia_charts(run_time_data_2, self.m_scatter_chart, self.m_scatter_series_2,
                                           self.m_pdf_chart, self.m_pdf_series_ls_2, self.m_pdf_series_map_2,
                                           self.m_cdf_chart, self.m_cdf_series_ls_2, self.m_cdf_series_map_2,
                                           self.m_relia_chart, self.m_relia_series_ls_2, self.m_relia_series_map_2,
                                           self.m_fail_chart, self.m_fail_series_ls_2, self.m_fail_series_map_2,
                                           self.ui.dev2ResultWidget)
                    self.run_time_data_1 = run_time_data_1
                    self.run_time_data_2 = run_time_data_2
                    # 比较分析
                    curr_time = max(np.max(run_time_data_1), np.max(run_time_data_2))
                    lambda_hat_1 = float(self.ui.dev1ResultWidget.item(Result.lamda_value.value).text())
                    beta_hat_1 = float(self.ui.dev1ResultWidget.item(Result.beta_value.value).text())
                    lambda_hat_2 = float(self.ui.dev2ResultWidget.item(Result.lamda_value.value).text())
                    beta_hat_2 = float(self.ui.dev2ResultWidget.item(Result.beta_value.value).text())
                    relia_dev_1 = rf.reliability(lambda_hat_1, beta_hat_1, curr_time)
                    relia_dev_2 = rf.reliability(lambda_hat_2, beta_hat_2, curr_time)
                    fail_dev_1 = rf.failure_rate(lambda_hat_1, beta_hat_1, curr_time)
                    fail_dev_2 = rf.failure_rate(lambda_hat_2, beta_hat_2, curr_time)
                    mtbf_dev_1 = float(self.ui.dev1ResultWidget.item(Result.mtbf_value.value).text())
                    mtbf_dev_2 = float(self.ui.dev2ResultWidget.item(Result.mtbf_value.value).text())
                    self.ui.cmpResultWidget.item(2).setText('%d' % curr_time)
                    if relia_dev_1 > relia_dev_2:
                        op = '>'
                    elif relia_dev_1 == relia_dev_2:
                        op = '='
                    else:
                        op = '<'
                    self.ui.cmpResultWidget.item(4).setText(
                        '设备1({:.3f}) {} 设备2({:.3f})'.format(relia_dev_1, op, relia_dev_2))
                    if fail_dev_1 > fail_dev_2:
                        op = '>'
                    elif fail_dev_1 == fail_dev_2:
                        op = '='
                    else:
                        op = '<'
                    self.ui.cmpResultWidget.item(6).setText(
                        '设备1({:.3f}) {} 设备2({:.3f})'.format(fail_dev_1, op, fail_dev_2))
                    if mtbf_dev_1 > mtbf_dev_2:
                        op = '>'
                    elif mtbf_dev_1 == mtbf_dev_2:
                        op = '='
                    else:
                        op = '<'
                    self.ui.cmpResultWidget.item(8).setText(
                        '设备1({:.3f}) {} 设备2({:.3f})'.format(mtbf_dev_1, op, mtbf_dev_2))
                else:
                    msg_box = QMessageBox(QMessageBox.Information,
                                          '提示',
                                          '设备均无故障',
                                          QMessageBox.Ok)
                    msg_box.exec_()
        else:
            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '请先选择设备！',
                                  QMessageBox.Ok)
            msg_box.exec_()

    # 可靠性分析页面按键事件
    @pyqtSlot()
    def on_rawDataButton_clicked(self):
        self.ui.reliaWidget.setCurrentIndex(0)

    @pyqtSlot()
    def on_scatterButton_clicked(self):
        self.ui.reliaWidget.setCurrentIndex(1)

    @pyqtSlot()
    def on_reliaButton_clicked(self):
        self.ui.reliaWidget.setCurrentIndex(2)

    def write_relia_result(self, filename):
        '''
        写入维修性分析结果
        :param filename: 文件名
        :return:
        '''
        item_head = ['设备名称', '尺度参数λ', '形状参数β', '可靠度', '失效率', 'MTBF']
        item_enum = [Result.lamda_value.value, Result.beta_value.value, Result.relia_value.value,
                     Result.fali_value.value, Result.mtbf_value.value]
        wb = Workbook()
        # 添加sheets
        relia_sheet = wb.create_sheet('可靠性分析结果', index=0)
        # 先写入表头
        relia_sheet.append(item_head)
        # 写入结果
        dev1_result = []
        dev2_result = []
        dev1_result.append(self.ui.devWidget.item(6).text())
        dev2_result.append(self.ui.devWidget_2.item(6).text())
        for i in item_enum:
            dev1_result.append(self.ui.dev1ResultWidget.item(i).text())
            dev2_result.append(self.ui.dev2ResultWidget.item(i).text())
        relia_sheet.append(dev1_result)
        relia_sheet.append(dev2_result)
        # 保存文档
        wb.save(filename)

    @pyqtSlot()
    def on_saveButton_clicked(self):
        save_dir = QFileDialog.getExistingDirectory(self, '选择文件夹', './')
        if save_dir != '':
            prefix = '可靠性分析_'
            # 保存可靠性分析图片
            if self.ui.reliaWidget.currentIndex() == 0:
                self.save_view(self.ui.tempView_1, os.path.join(save_dir, '环境温度折线图.jpg'))
                self.save_view(self.ui.tempView_2, os.path.join(save_dir, '刀头温度折线图.jpg'))
                self.save_view(self.ui.rpaView, os.path.join(save_dir, '重复定位精度折线图.jpg'))
                self.save_view(self.ui.straView, os.path.join(save_dir, '温差折线图.jpg'))
            elif self.ui.reliaWidget.currentIndex() == 1:
                self.save_view(self.ui.scatterView, os.path.join(save_dir, '故障时间间隔散点图.jpg'))
            else:
                self.save_view(self.ui.pdfView, os.path.join(save_dir, '威布尔分布概率密度图.jpg'))
                self.save_view(self.ui.cdfView, os.path.join(save_dir, '威布尔分布累积分布图.jpg'))
                self.save_view(self.ui.reliaView, os.path.join(save_dir, '可靠度曲线图.jpg'))
                self.save_view(self.ui.failView, os.path.join(save_dir, '失效率曲线图.jpg'))
                # 将结果写入表格
                self.write_relia_result(os.path.join(save_dir, '可靠性分析结果.xlsx'))

            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '保存成功',
                                  QMessageBox.Ok)
            msg_box.exec_()

    @pyqtSlot()
    def on_outputDataButton_clicked(self):
        item_head = ['运行时间/h', '环境温度/℃', '刀头温度/℃', '重复定位精度/μm', '设备名称', '记录时间']
        wb = Workbook()
        # 添加sheets
        x_sheet = wb.create_sheet('X轴', index=0)
        y_sheet = wb.create_sheet('Y轴', index=1)
        z_sheet = wb.create_sheet('Z轴', index=2)
        # 先写入表头
        x_sheet.append(item_head)
        y_sheet.append(item_head)
        z_sheet.append(item_head)
        sql_str = 'select r.run_time, r.env_temp, r.kni_temp, r.rpa, r.axis, ' \
                  '(select d.name from device d where d.id = r.dev_id), r.record_time from record r'
        self.query.prepare(sql_str)
        if not self.query.exec_():
            print(self.query.lastError().text())
        else:
            while self.query.next():
                # 写入结果
                axis = self.query.value(4)
                if 'X' == axis:
                    item = [self.query.value(0), self.query.value(1),
                            self.query.value(2), self.query.value(3),
                            self.query.value(5), self.query.value(6)]
                    x_sheet.append(item)
                elif 'Y' == axis:
                    item = [self.query.value(0), self.query.value(1),
                            self.query.value(2), self.query.value(3),
                            self.query.value(5), self.query.value(6)]
                    y_sheet.append(item)
                else:
                    item = [self.query.value(0), self.query.value(1),
                            self.query.value(2), self.query.value(3),
                            self.query.value(5), self.query.value(6)]
                    z_sheet.append(item)
        filename, _ = QFileDialog.getSaveFileName(self, '保存文件', './', 'Excel 2007 (*.xlsx);;Excel 2003 (*.xls)')
        if filename.strip(' ') != '':
            wb.save(filename)
            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '保存成功',
                                  QMessageBox.Ok)
            msg_box.exec_()

    @pyqtSlot()
    def on_manaDevButton_clicked(self):
        dev_widget = DeviceListWidget()
        dev_widget.show()
        dev_widget.exec_()
        dev_widget.destroy()
        self.initDevInfo()

    @pyqtSlot()
    def on_manaRecButton_clicked(self):
        rec_widget = RecordListDialog()
        rec_widget.show()
        rec_widget.exec_()
        rec_widget.destroy()

    def read_maintain_data(self, dev_id):
        '''
        读取维修数据
        :return:
        '''
        maintain_times = []  # 保存读取数据
        # 数据库查询
        sql_str = 'select maintain_time from maintain where dev_id = :dev_id'
        self.query.prepare(sql_str)
        self.query.bindValue(':dev_id', dev_id)
        if not self.query.exec_():
            print(self.query.lastError().text())
        else:
            while self.query.next():
                maintain_times.append(self.query.value(0))
        return maintain_times

    # 维修性分析页面按键事件
    def plot_maintain_charts(self, maintain_time_data, m_scatter_chart_maintain, m_scatter_series_maintain,
                             m_pdf_chart_maintain, m_pdf_series_maintain_ls, m_pdf_series_maintain_map,
                             m_cdf_chart_maintain, m_cdf_series_maintain_ls, m_cdf_series_maintain_map,
                             m_relia_chart_maintain, m_relia_series_maintain_ls, m_relia_series_maintain_map,
                             m_fail_chart_maintain, m_fail_series_maintain_ls, m_fail_series_maintain_map,
                             resultWidget):
        maintain_time_data = np.array(maintain_time_data)
        t_max = max(maintain_time_data)
        t_range = np.arange(1, t_max, 1)
        maintain_time_prob = np.array(
            [(i - 0.3) / (len(maintain_time_data) + 0.4) for i in range(1, len(maintain_time_data) + 1)])
        # 绘制散点图
        m_scatter_series_maintain.clear()
        for t, prob in zip(maintain_time_data, maintain_time_prob):
            m_scatter_series_maintain.append(t, prob)
        # 重新设定x轴范围
        m_scatter_chart_maintain.axisX().setRange(0, max(t_max, m_scatter_chart_maintain.axisX().max()))
        m_pdf_chart_maintain.axisX().setRange(0, max(t_max, m_pdf_chart_maintain.axisX().max()))
        m_cdf_chart_maintain.axisX().setRange(0, max(t_max, m_cdf_chart_maintain.axisX().max()))
        m_relia_chart_maintain.axisX().setRange(0, max(t_max, m_relia_chart_maintain.axisX().max()))
        m_fail_chart_maintain.axisX().setRange(0, max(t_max, m_fail_chart_maintain.axisX().max()))
        ## 以下是最小二乘法计算
        lambda_hat_mt, beta_hat_mt, pdf_prob_ls_mt, cdf_prob_ls_mt, relia_prob_ls_mt, fail_prob_ls_mt = self.compute_curves(
            maintain_time_data, method=0)
        lambda_map_mt, beta_map_mt, pdf_prob_map_mt, cdf_prob_map_mt, relia_prob_map_mt, fail_prob_map_mt = self.compute_curves(
            maintain_time_data, method=1)
        # 保存全局变量
        if m_pdf_series_maintain_ls == self.m_pdf_series_maintain_ls:
            self.lambda_hat_mt_1, self.beta_hat_mt_1 = lambda_hat_mt, beta_hat_mt
        else:
            self.lambda_hat_mt_2, self.beta_hat_mt_2 = lambda_hat_mt, beta_hat_mt
        if m_pdf_series_maintain_map == self.m_pdf_series_maintain_map:
            self.lambda_map_mt_1, self.beta_map_mt_1 = lambda_map_mt, beta_map_mt
        else:
            self.lambda_map_mt_2, self.beta_map_mt_2 = lambda_map_mt, beta_map_mt
        # 重新设定y轴范围
        m_pdf_chart_maintain.axisY().setRange(0, max(np.max(pdf_prob_ls_mt), np.max(pdf_prob_map_mt),
                                                     m_pdf_chart_maintain.axisY().max()))
        m_cdf_chart_maintain.axisY().setRange(0, max(np.max(cdf_prob_ls_mt), np.max(cdf_prob_map_mt),
                                                     m_cdf_chart_maintain.axisY().max()))
        m_relia_chart_maintain.axisY().setRange(0, max(np.max(relia_prob_ls_mt), np.max(relia_prob_map_mt),
                                                       m_relia_chart_maintain.axisY().max()))
        m_fail_chart_maintain.axisY().setRange(0, max(np.max(fail_prob_ls_mt), np.max(fail_prob_map_mt),
                                                      m_fail_chart_maintain.axisY().max()))
        method = self.ui.comboBox_2.currentIndex()  # 获取算法选择索引,0---最小二乘法, 1---贝叶斯估计法
        if method == 0:
            m_pdf_chart_maintain.removeSeries(self.m_pdf_series_maintain_map)
            m_cdf_chart_maintain.removeSeries(self.m_cdf_series_maintain_map)
            m_relia_chart_maintain.removeSeries(self.m_relia_series_maintain_map)
            m_fail_chart_maintain.removeSeries(self.m_fail_series_maintain_map)

            m_pdf_chart_maintain.removeSeries(self.m_pdf_series_maintain_map_2)
            m_cdf_chart_maintain.removeSeries(self.m_cdf_series_maintain_map_2)
            m_relia_chart_maintain.removeSeries(self.m_relia_series_maintain_map_2)
            m_fail_chart_maintain.removeSeries(self.m_fail_series_maintain_map_2)
            # 填充表格
            resultWidget.item(Result.lamda_value.value).setText('%.4f' % lambda_hat_mt)
            resultWidget.item(Result.beta_value.value).setText('%.4f' % beta_hat_mt)
            resultWidget.item(Result.mtbf_value.value).setText('%.4f' % rf.mtbf(lambda_hat_mt, beta_hat_mt))
        else:
            m_pdf_chart_maintain.removeSeries(self.m_pdf_series_maintain_ls)
            m_cdf_chart_maintain.removeSeries(self.m_cdf_series_maintain_ls)
            m_relia_chart_maintain.removeSeries(self.m_relia_series_maintain_ls)
            m_fail_chart_maintain.removeSeries(self.m_fail_series_maintain_ls)

            m_pdf_chart_maintain.removeSeries(self.m_pdf_series_maintain_ls_2)
            m_cdf_chart_maintain.removeSeries(self.m_cdf_series_maintain_ls_2)
            m_relia_chart_maintain.removeSeries(self.m_relia_series_maintain_ls_2)
            m_fail_chart_maintain.removeSeries(self.m_fail_series_maintain_ls_2)
            # 填充表格
            resultWidget.item(Result.lamda_value.value).setText('%.4f' % lambda_map_mt)
            resultWidget.item(Result.beta_value.value).setText('%.4f' % beta_map_mt)
            resultWidget.item(Result.mtbf_value.value).setText(
                '%.4f' % rf.mtbf(lambda_map_mt, beta_map_mt))
        # 绘制概率密度曲线
        m_pdf_series_maintain_ls.clear()
        for t, prob in zip(t_range, pdf_prob_ls_mt):
            m_pdf_series_maintain_ls.append(t, prob)
        # 绘制累积分布曲线
        m_cdf_series_maintain_ls.clear()
        for t, prob in zip(t_range, cdf_prob_ls_mt):
            m_cdf_series_maintain_ls.append(t, prob)
        # 绘制可靠度曲线
        m_relia_series_maintain_ls.clear()
        for t, prob in zip(t_range, relia_prob_ls_mt):
            m_relia_series_maintain_ls.append(t, prob)
        # 绘制失效率曲线
        m_fail_series_maintain_ls.clear()
        for t, prob in zip(t_range, fail_prob_ls_mt):
            m_fail_series_maintain_ls.append(t, prob)
        # 绘制概率密度曲线
        m_pdf_series_maintain_map.clear()
        for t, prob in zip(t_range, pdf_prob_map_mt):
            m_pdf_series_maintain_map.append(t, prob)
        # 绘制累积分布曲线
        m_cdf_series_maintain_map.clear()
        for t, prob in zip(t_range, cdf_prob_map_mt):
            m_cdf_series_maintain_map.append(t, prob)
        # 绘制可靠度曲线
        m_relia_series_maintain_map.clear()
        for t, prob in zip(t_range, relia_prob_map_mt):
            m_relia_series_maintain_map.append(t, prob)
        # 绘制失效率曲线
        m_fail_series_maintain_map.clear()
        for t, prob in zip(t_range, fail_prob_map_mt):
            m_fail_series_maintain_map.append(t, prob)

    @pyqtSlot()
    def on_loadButton_2_clicked(self):
        dev_1 = self.ui.devWidget.item(2).text()
        dev_2 = self.ui.devWidget_2.item(2).text()
        if dev_1.strip(' ') != '' and dev_2.strip(' ') == '':  # 只有设备1的数据
            dev_1 = int(dev_1)
            maintain_time_of_dev_1 = self.read_maintain_data(dev_1)
            if len(maintain_time_of_dev_1) <= 0:
                self.maintain_time_of_dev_1 = maintain_time_of_dev_1
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '无维修数据',
                                      QMessageBox.Ok)
                msg_box.exec_()
                return
            else:
                # ---维修性分析相关---
                self.plot_maintain_charts(maintain_time_of_dev_1, self.m_scatter_chart_maintain,
                                          self.m_scatter_series_maintain,
                                          self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls,
                                          self.m_pdf_series_maintain_map,
                                          self.m_cdf_chart_maintain, self.m_cdf_series_maintain_ls,
                                          self.m_cdf_series_maintain_map,
                                          self.m_relia_chart_maintain, self.m_relia_series_maintain_ls,
                                          self.m_relia_series_maintain_map,
                                          self.m_fail_chart_maintain, self.m_fail_series_maintain_ls,
                                          self.m_fail_series_maintain_map,
                                          self.ui.dev1MtResultWidget)
        elif dev_1.strip(' ') == '' and dev_2.strip(' ') != '':  # 只有设备2的数据
            dev_2 = int(dev_2)
            maintain_time_of_dev_2 = self.read_maintain_data(dev_2)
            if len(maintain_time_of_dev_2) <= 0:
                self.maintain_time_of_dev_2 = maintain_time_of_dev_2
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '无维修数据',
                                      QMessageBox.Ok)
                msg_box.exec_()
                return
            else:
                # ---维修性分析相关---
                self.plot_maintain_charts(maintain_time_of_dev_2, self.m_scatter_chart_maintain,
                                          self.m_scatter_series_maintain_2,
                                          self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls_2,
                                          self.m_pdf_series_maintain_map_2,
                                          self.m_cdf_chart_maintain, self.m_cdf_series_maintain_ls_2,
                                          self.m_cdf_series_maintain_map_2,
                                          self.m_relia_chart_maintain, self.m_relia_series_maintain_ls_2,
                                          self.m_relia_series_maintain_map_2,
                                          self.m_fail_chart_maintain, self.m_fail_series_maintain_ls_2,
                                          self.m_fail_series_maintain_map_2,
                                          self.ui.dev2MtResultWidget)
        elif dev_1.strip(' ') != '' and dev_2.strip(' ') != '':
            dev_1 = int(dev_1)
            dev_2 = int(dev_2)
            maintain_time_of_dev_1 = self.read_maintain_data(dev_1)
            maintain_time_of_dev_2 = self.read_maintain_data(dev_2)
            if len(maintain_time_of_dev_1) <= 0 and len(maintain_time_of_dev_2) > 0:  # 只有设备2有数据
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '设备1无维修数据',
                                      QMessageBox.Ok)
                msg_box.exec_()
                self.plot_maintain_charts(maintain_time_of_dev_2, self.m_scatter_chart_maintain,
                                          self.m_scatter_series_maintain_2,
                                          self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls_2,
                                          self.m_pdf_series_maintain_map_2,
                                          self.m_cdf_chart_maintain, self.m_cdf_series_maintain_ls_2,
                                          self.m_cdf_series_maintain_map_2,
                                          self.m_relia_chart_maintain, self.m_relia_series_maintain_ls_2,
                                          self.m_relia_series_maintain_map_2,
                                          self.m_fail_chart_maintain, self.m_fail_series_maintain_ls_2,
                                          self.m_fail_series_maintain_map_2,
                                          self.ui.dev2MtResultWidget)
                return
            elif len(maintain_time_of_dev_1) > 0 and len(maintain_time_of_dev_2) <= 0:  # 只有设备1有数据:
                # ---维修性分析相关---
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '设备2无维修数据',
                                      QMessageBox.Ok)
                msg_box.exec_()
                self.plot_maintain_charts(maintain_time_of_dev_1, self.m_scatter_chart_maintain,
                                          self.m_scatter_series_maintain,
                                          self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls,
                                          self.m_pdf_series_maintain_map,
                                          self.m_cdf_chart_maintain, self.m_cdf_series_maintain_ls,
                                          self.m_cdf_series_maintain_map,
                                          self.m_relia_chart_maintain, self.m_relia_series_maintain_ls,
                                          self.m_relia_series_maintain_map,
                                          self.m_fail_chart_maintain, self.m_fail_series_maintain_ls,
                                          self.m_fail_series_maintain_map,
                                          self.ui.dev1MtResultWidget)
            elif len(maintain_time_of_dev_1) > 0 and len(maintain_time_of_dev_2) > 0:
                self.maintain_time_of_dev_1 = maintain_time_of_dev_1
                self.maintain_time_of_dev_2 = maintain_time_of_dev_2
                self.plot_maintain_charts(maintain_time_of_dev_1, self.m_scatter_chart_maintain,
                                          self.m_scatter_series_maintain,
                                          self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls,
                                          self.m_pdf_series_maintain_map,
                                          self.m_cdf_chart_maintain, self.m_cdf_series_maintain_ls,
                                          self.m_cdf_series_maintain_map,
                                          self.m_relia_chart_maintain, self.m_relia_series_maintain_ls,
                                          self.m_relia_series_maintain_map,
                                          self.m_fail_chart_maintain, self.m_fail_series_maintain_ls,
                                          self.m_fail_series_maintain_map,
                                          self.ui.dev1MtResultWidget)
                self.plot_maintain_charts(maintain_time_of_dev_2, self.m_scatter_chart_maintain,
                                          self.m_scatter_series_maintain_2,
                                          self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls_2,
                                          self.m_pdf_series_maintain_map_2,
                                          self.m_cdf_chart_maintain, self.m_cdf_series_maintain_ls_2,
                                          self.m_cdf_series_maintain_map_2,
                                          self.m_relia_chart_maintain, self.m_relia_series_maintain_ls_2,
                                          self.m_relia_series_maintain_map_2,
                                          self.m_fail_chart_maintain, self.m_fail_series_maintain_ls_2,
                                          self.m_fail_series_maintain_map_2,
                                          self.ui.dev2MtResultWidget)
                # 比较分析
                curr_time = max(np.max(maintain_time_of_dev_1), np.max(maintain_time_of_dev_2))
                lambda_hat_1 = float(self.ui.dev1MtResultWidget.item(Result.lamda_value.value).text())
                beta_hat_1 = float(self.ui.dev1MtResultWidget.item(Result.beta_value.value).text())
                lambda_hat_2 = float(self.ui.dev2MtResultWidget.item(Result.lamda_value.value).text())
                beta_hat_2 = float(self.ui.dev2MtResultWidget.item(Result.beta_value.value).text())
                relia_dev_1 = rf.reliability(lambda_hat_1, beta_hat_1, curr_time)
                relia_dev_2 = rf.reliability(lambda_hat_2, beta_hat_2, curr_time)
                fail_dev_1 = rf.failure_rate(lambda_hat_1, beta_hat_1, curr_time)
                fail_dev_2 = rf.failure_rate(lambda_hat_2, beta_hat_2, curr_time)
                mtbf_dev_1 = float(self.ui.dev1MtResultWidget.item(Result.mtbf_value.value).text())
                mtbf_dev_2 = float(self.ui.dev2MtResultWidget.item(Result.mtbf_value.value).text())
                self.ui.cmpMtResultWidget.item(2).setText('%d' % curr_time)
                if relia_dev_1 > relia_dev_2:
                    op = '>'
                elif relia_dev_1 == relia_dev_2:
                    op = '='
                else:
                    op = '<'
                self.ui.cmpMtResultWidget.item(4).setText(
                    '设备1({:.3f}) {} 设备2({:.3f})'.format(relia_dev_1, op, relia_dev_2))
                if fail_dev_1 > fail_dev_2:
                    op = '>'
                elif fail_dev_1 == fail_dev_2:
                    op = '='
                else:
                    op = '<'
                self.ui.cmpMtResultWidget.item(6).setText(
                    '设备1({:.3f}) {} 设备2({:.3f})'.format(fail_dev_1, op, fail_dev_2))
                if mtbf_dev_1 > mtbf_dev_2:
                    op = '>'
                elif mtbf_dev_1 == mtbf_dev_2:
                    op = '='
                else:
                    op = '<'
                self.ui.cmpMtResultWidget.item(8).setText(
                    '设备1({:.3f}) {} 设备2({:.3f})'.format(mtbf_dev_1, op, mtbf_dev_2))
            else:
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '均无维修数据',
                                      QMessageBox.Ok)
                msg_box.exec_()
        else:
            pass

    @pyqtSlot()
    def on_nextMtResultButton_clicked(self):
        page_num = self.ui.mtResultWidget.count()
        index = self.ui.mtResultWidget.currentIndex()
        index += 1
        if index == page_num:
            index = 0
        self.ui.mtResultWidget.setCurrentIndex(index)

    @pyqtSlot()
    def on_scatterButton_2_clicked(self):
        self.ui.maintainWidget.setCurrentIndex(0)

    @pyqtSlot()
    def on_reliaButton_2_clicked(self):
        self.ui.maintainWidget.setCurrentIndex(1)

    def write_maintain_result(self, filename):
        '''
        写入维修性分析结果
        :param filename: 文件名
        :return:
        '''
        item_head = ['设备名称', '尺度参数λ', '形状参数β', '修复率', '失效率', 'MTTR']
        item_enum = [Result.lamda_value.value, Result.beta_value.value, Result.relia_value.value,
                     Result.fali_value.value, Result.mtbf_value.value]
        wb = Workbook()
        # 添加sheets
        maintain_sheet = wb.create_sheet('维修性分析结果', index=0)
        # 先写入表头
        maintain_sheet.append(item_head)
        # 写入结果
        dev1_result = []
        dev2_result = []
        dev1_result.append(self.ui.devWidget.item(6).text())
        dev2_result.append(self.ui.devWidget_2.item(6).text())
        for i in item_enum:
            dev1_result.append(self.ui.dev1MtResultWidget.item(i).text())
            dev2_result.append(self.ui.dev2MtResultWidget.item(i).text())
        maintain_sheet.append(dev1_result)
        maintain_sheet.append(dev2_result)
        # 保存文档
        wb.save(filename)

    @pyqtSlot()
    def on_saveButton_2_clicked(self):
        save_dir = QFileDialog.getExistingDirectory(self, '选择文件夹', './')
        if save_dir != '':
            prefix = '维修性分析_'
            # 保存维修性分析图片
            if self.ui.maintainWidget.currentIndex() == 0:
                self.save_view(self.ui.scatterView_2, os.path.join(save_dir, '维修时间散点图.jpg'))
            else:
                self.save_view(self.ui.pdfView_2, os.path.join(save_dir, '威布尔分布概率密度图.jpg'))
                self.save_view(self.ui.cdfView_2, os.path.join(save_dir, '威布尔分布累积分布图.jpg'))
                self.save_view(self.ui.reliaView_2, os.path.join(save_dir, '维修度曲线图.jpg'))
                self.save_view(self.ui.failView_2, os.path.join(save_dir, '失效率曲线图.jpg'))
                # 将结果写入表格
                self.write_maintain_result(os.path.join(save_dir, '维修性分析结果.xlsx'))
            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '保存成功',
                                  QMessageBox.Ok)
            msg_box.exec_()

    @pyqtSlot()
    def on_outputDataButton_2_clicked(self):
        item_head = ['序号', '维修日期', '维修人员', '维修时长/h', '设备名称', '记录时间']
        wb = Workbook()
        # 添加sheets
        maintain_sheet = wb.create_sheet('维修记录', index=0)
        # 先写入表头
        maintain_sheet.append(item_head)
        sql_str = 'select m.id, m.maintain_date, m.person, m.maintain_time, ' \
                  '(select d.name from device d where d.id = m.dev_id), ' \
                  'm.record_time from maintain m'
        self.query.prepare(sql_str)
        if not self.query.exec_():
            print(self.query.lastError().text())
        else:
            while self.query.next():
                # 写入结果
                maintain_sheet.append([self.query.value(i) for i in range(0, len(item_head))])
        filename, _ = QFileDialog.getSaveFileName(self, '保存文件', './', 'Excel 2007 (*.xlsx);;Excel 2003 (*.xls)')
        if filename.strip(' ') != '':
            wb.save(filename)
            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '保存成功',
                                  QMessageBox.Ok)
            msg_box.exec_()

    @pyqtSlot()
    def on_manaMtRecButton_clicked(self):
        mt_widget = MaintainListDialog()
        mt_widget.show()
        mt_widget.exec_()
        mt_widget.destroy()

    # 故障分析页面响应事件
    def read_fault_data(self, dev_id):
        '''
        读取故障数据
        :param filename: 文件名
        :return:
        '''
        item_num = 4
        data_table = dict()
        # 数据库查询
        # 整机
        whole_table = []
        patts, posis, reasons, roots, statuses = [], [], [], [], []
        query_sql = 'select pattern, position, reason, root, status from breakdown ' \
                    'where dev_id = :dev_id and status = :status'
        self.query.prepare(query_sql)
        self.query.bindValue(':dev_id', dev_id)
        self.query.bindValue(':status', 1)
        if not self.query.exec_():
            print(self.query.lastError().text())
            return data_table
        else:
            while self.query.next():
                patts.append(str(self.query.value(0)))
                posis.append(str(self.query.value(1)))
                reasons.append(str(self.query.value(2)))
                roots.append(str(self.query.value(3)))
                statuses.append(int(self.query.value(4)))
            whole_table.append(patts)
            whole_table.append(posis)
            whole_table.append(reasons)
            whole_table.append(roots)
            whole_table.append(statuses)
            # 整机
            subsys_table = []
            patts, posis, reasons, roots, statuses = [], [], [], [], []
            query_sql = 'select pattern, position, reason, root, status from breakdown ' \
                        'where dev_id = :dev_id and status = :status'
            self.query.prepare(query_sql)
            self.query.bindValue(':dev_id', dev_id)
            self.query.bindValue(':status', 0)
            if not self.query.exec_():
                print(self.query.lastError().text())
                return data_table
            else:
                while self.query.next():
                    patts.append(str(self.query.value(0)))
                    posis.append(str(self.query.value(1)))
                    reasons.append(str(self.query.value(2)))
                    roots.append(str(self.query.value(3)))
                    statuses.append(int(self.query.value(4)))
                subsys_table.append(patts)
                subsys_table.append(posis)
                subsys_table.append(reasons)
                subsys_table.append(roots)
                subsys_table.append(statuses)
            data_table['whole'] = whole_table
            data_table['subsys'] = subsys_table
        return data_table

    def count_fault_data(self, data_table):
        # 创建各故障数据项集合
        patt_set = dict()
        for rec in data_table[PATT]:
            if rec in patt_set:
                patt_set[rec] += 1 / len(data_table[PATT])
            else:
                patt_set[rec] = 0
        patt_set = dict(sorted(patt_set.items(), key=operator.itemgetter(0)))
        posi_set = dict()
        for rec in data_table[POSI]:
            if rec in posi_set:
                posi_set[rec] += 1 / len(data_table[POSI])
            else:
                posi_set[rec] = 0
        posi_set = dict(sorted(posi_set.items(), key=operator.itemgetter(0)))
        reason_set = dict()
        for rec in data_table[REASON]:
            if rec in reason_set:
                reason_set[rec] += 1 / len(data_table[REASON])
            else:
                reason_set[rec] = 0
        reason_set = dict(sorted(reason_set.items(), key=operator.itemgetter(0)))
        root_set = dict()
        for rec in data_table[ROOT]:
            if rec in root_set:
                root_set[rec] += 1 / len(data_table[ROOT])
            else:
                root_set[rec] = 0
        root_set = dict(sorted(root_set.items(), key=operator.itemgetter(0)))
        fault_set = [patt_set, posi_set, reason_set, root_set]
        return fault_set

    def plot_fault_bar(self, fault_datas, dev_name='设备'):
        '''
        绘制故障统计图
        :param fault_data:
        :return:
        '''
        # 记录统计结果的字符串
        whole_patt = ''
        whole_posi = ''
        whole_reason = ''
        whole_root = ''
        subsys_patt = ''
        subsys_posi = ''
        subsys_reason = ''
        subsys_root = ''
        # 清除旧记录
        self.m_patt_series.clear()
        self.m_posi_series.clear()
        self.m_reason_series.clear()
        self.m_root_series.clear()
        self.m_patt_series_2.clear()
        self.m_posi_series_2.clear()
        self.m_reason_series_2.clear()
        self.m_root_series_2.clear()
        self.m_patt_chart.removeSeries(self.m_patt_series)
        self.m_posi_chart.removeSeries(self.m_posi_series)
        self.m_reason_chart.removeSeries(self.m_reason_series)
        self.m_root_chart.removeSeries(self.m_root_series)
        self.m_patt_chart_2.removeSeries(self.m_patt_series_2)
        self.m_posi_chart_2.removeSeries(self.m_posi_series_2)
        self.m_reason_chart_2.removeSeries(self.m_reason_series_2)
        self.m_root_chart_2.removeSeries(self.m_root_series_2)
        ## 整机
        for i, fault_data in enumerate(fault_datas):
            whole_data = fault_data['whole']
            fault_set = self.count_fault_data(whole_data)
            # --故障模式--
            patt_set = fault_set[PATT]
            # 定义柱状条
            # 柱状条赋值
            patt_items = []
            patt_probs = []
            bar_set = QBarSet(dev_name + str(i + 1))
            for key, value in patt_set.items():
                patt_items.append(key)
                patt_probs.append(value)
                bar_set.append(value)
            # 定义横坐标轴
            if i == 0:
                axisX = QBarCategoryAxis()
                axisX.setTitleFont(QFont('Times New Roman', 10))  # 设置横坐标标题字体的类型和大小
                axisX.setLabelsFont(QFont('Times New Roman', 10))  # 设置横坐标刻度的字体类型和大小
                axisX.append(patt_items)
                # 设置柱状集
                self.m_patt_chart.setAxisX(axisX)
            self.m_patt_chart.axisY().setRange(0, max(max(patt_probs), self.m_patt_chart.axisY().max()))
            self.m_patt_series.append(bar_set)
            self.m_patt_chart.legend().setVisible(True)
            self.m_patt_chart.legend().setAlignment(Qt.AlignBottom)
            # --故障部位--
            posi_set = fault_set[POSI]
            # 定义柱状条
            # 柱状条赋值
            posi_items = []
            posi_probs = []
            bar_set = QBarSet(dev_name + str(i + 1))
            for key, value in posi_set.items():
                posi_items.append(key)
                posi_probs.append(value)
                bar_set.append(value)
            # 定义横坐标轴
            if i == 0:
                axisX = QBarCategoryAxis()
                axisX.setTitleFont(QFont('Times New Roman', 10))  # 设置横坐标标题字体的类型和大小
                axisX.setLabelsFont(QFont('Times New Roman', 10))  # 设置横坐标刻度的字体类型和大小
                axisX.append(posi_items)
                # 设置柱状集
                self.m_posi_chart.setAxisX(axisX)
            self.m_posi_chart.axisY().setRange(0, max(max(posi_probs), self.m_posi_chart.axisY().max()))
            self.m_posi_series.append(bar_set)
            self.m_posi_chart.legend().setVisible(True)
            self.m_posi_chart.legend().setAlignment(Qt.AlignBottom)
            # --故障原因-
            reason_set = fault_set[REASON]
            # 定义柱状条
            # 柱状条赋值
            reason_items = []
            reason_probs = []
            bar_set = QBarSet(dev_name + str(i + 1))
            for key, value in reason_set.items():
                reason_items.append(key)
                reason_probs.append(value)
                bar_set.append(value)
            # 定义横坐标轴
            if i == 0:
                axisX = QBarCategoryAxis()
                axisX.setTitleFont(QFont('Times New Roman', 10))  # 设置横坐标标题字体的类型和大小
                axisX.setLabelsFont(QFont('Times New Roman', 10))  # 设置横坐标刻度的字体类型和大小
                axisX.append(reason_items)
                # 设置柱状集
                self.m_reason_chart.setAxisX(axisX)
            self.m_reason_chart.axisY().setRange(0, max(max(reason_probs), self.m_reason_chart.axisY().max()))
            self.m_reason_series.append(bar_set)
            self.m_reason_chart.legend().setVisible(True)
            self.m_reason_chart.legend().setAlignment(Qt.AlignBottom)
            # --故障溯源--
            root_set = fault_set[ROOT]
            # 定义柱状条
            # 柱状条赋值
            root_items = []
            root_probs = []
            bar_set = QBarSet(dev_name + str(i + 1))
            for key, value in root_set.items():
                root_items.append(key)
                root_probs.append(value)
                bar_set.append(value)
            # 定义横坐标轴
            if i == 0:
                axisX = QBarCategoryAxis()
                axisX.setTitleFont(QFont('Times New Roman', 10))  # 设置横坐标标题字体的类型和大小
                axisX.setLabelsFont(QFont('Times New Roman', 10))  # 设置横坐标刻度的字体类型和大小
                axisX.append(root_items)
                # 设置柱状集
                self.m_root_chart.setAxisX(axisX)
            self.m_root_chart.axisY().setRange(0, max(max(root_probs), self.m_root_chart.axisY().max()))
            self.m_root_series.append(bar_set)
            self.m_root_chart.legend().setVisible(True)
            self.m_root_chart.legend().setAlignment(Qt.AlignBottom)
            # 保存结果到字符串
            whole_patt += dev_name + str(i + 1) + ': ' + patt_items[patt_probs.index(max(patt_probs))] + ' '
            whole_posi += dev_name + str(i + 1) + ': ' + posi_items[posi_probs.index(max(posi_probs))] + ' '
            whole_reason += dev_name + str(i + 1) + ': ' + reason_items[reason_probs.index(max(reason_probs))] + ' '
            whole_root += dev_name + str(i + 1) + ': ' + root_items[root_probs.index(max(root_probs))] + ' '
            ##--子系统--
            subsys_data = fault_data['subsys']
            # 统计各故障频率
            fault_set = self.count_fault_data(subsys_data)
            # --故障模式--
            patt_set = fault_set[PATT]
            # 定义柱状条
            # 柱状条赋值
            patt_items = []
            patt_probs = []
            bar_set = QBarSet(dev_name + str(i + 1))
            for key, value in patt_set.items():
                patt_items.append(key)
                patt_probs.append(value)
                bar_set.append(value)
            # 定义横坐标轴
            if i == 0:
                axisX = QBarCategoryAxis()
                axisX.setTitleFont(QFont('Times New Roman', 10))  # 设置横坐标标题字体的类型和大小
                axisX.setLabelsFont(QFont('Times New Roman', 10))  # 设置横坐标刻度的字体类型和大小
                axisX.append(patt_set)
                # 设置柱状集
                self.m_patt_chart_2.setAxisX(axisX)
            self.m_patt_chart_2.axisY().setRange(0, max(max(patt_probs), self.m_patt_chart_2.axisY().max()))
            self.m_patt_series_2.append(bar_set)
            self.m_patt_chart_2.legend().setVisible(True)
            self.m_patt_chart_2.legend().setAlignment(Qt.AlignBottom)
            # --故障部位--
            posi_set = fault_set[POSI]
            # 定义柱状条
            # 柱状条赋值
            posi_items = []
            posi_probs = []
            bar_set = QBarSet(dev_name + str(i + 1))
            for key, value in posi_set.items():
                posi_items.append(key)
                posi_probs.append(value)
                bar_set.append(value)
            # 定义横坐标轴
            if i == 0:
                axisX = QBarCategoryAxis()
                axisX.setTitleFont(QFont('Times New Roman', 10))  # 设置横坐标标题字体的类型和大小
                axisX.setLabelsFont(QFont('Times New Roman', 10))  # 设置横坐标刻度的字体类型和大小
                axisX.append(posi_set)
                # 设置柱状集
                self.m_posi_chart_2.setAxisX(axisX)
            self.m_posi_chart_2.axisY().setRange(0, max(max(posi_probs), self.m_posi_chart_2.axisY().max()))
            self.m_posi_series_2.append(bar_set)
            self.m_posi_chart_2.legend().setVisible(True)
            self.m_posi_chart_2.legend().setAlignment(Qt.AlignBottom)
            # --故障原因-
            reason_set = fault_set[REASON]
            # 定义柱状条
            # 柱状条赋值
            reason_items = []
            reason_probs = []
            bar_set = QBarSet(dev_name + str(i + 1))
            for key, value in reason_set.items():
                reason_items.append(key)
                reason_probs.append(value)
                bar_set.append(value)
            # 定义横坐标轴
            if i == 0:
                axisX = QBarCategoryAxis()
                axisX.setTitleFont(QFont('Times New Roman', 10))  # 设置横坐标标题字体的类型和大小
                axisX.setLabelsFont(QFont('Times New Roman', 10))  # 设置横坐标刻度的字体类型和大小
                axisX.append(reason_set)
                # 设置柱状集
                self.m_reason_chart_2.setAxisX(axisX)
            self.m_reason_chart_2.axisY().setRange(0, max(max(reason_probs), self.m_reason_chart_2.axisY().max()))
            self.m_reason_series_2.append(bar_set)
            self.m_reason_chart_2.legend().setVisible(True)
            self.m_reason_chart_2.legend().setAlignment(Qt.AlignBottom)
            # --故障溯源--
            root_set = fault_set[ROOT]
            # 定义柱状条
            # 柱状条赋值
            root_items = []
            root_probs = []
            bar_set = QBarSet(dev_name + str(i + 1))
            for key, value in root_set.items():
                root_items.append(key)
                root_probs.append(value)
                bar_set.append(value)
            # 定义横坐标轴
            if i == 0:
                axisX = QBarCategoryAxis()
                axisX.setTitleFont(QFont('Times New Roman', 10))  # 设置横坐标标题字体的类型和大小
                axisX.setLabelsFont(QFont('Times New Roman', 10))  # 设置横坐标刻度的字体类型和大小
                axisX.append(root_set)
                # 设置柱状集
                self.m_root_chart_2.setAxisX(axisX)
            self.m_root_chart_2.axisY().setRange(0, max(max(root_probs), self.m_root_chart_2.axisY().max()))
            self.m_root_series_2.append(bar_set)
            self.m_root_chart_2.legend().setVisible(True)
            self.m_root_chart_2.legend().setAlignment(Qt.AlignBottom)
            # 保存结果到字符串
            subsys_patt += dev_name + str(i + 1) + ': ' + patt_items[patt_probs.index(max(patt_probs))] + ' '
            subsys_posi += dev_name + str(i + 1) + ': ' + posi_items[posi_probs.index(max(posi_probs))] + ' '
            subsys_reason += dev_name + str(i + 1) + ': ' + reason_items[reason_probs.index(max(reason_probs))] + ' '
            subsys_root += dev_name + str(i + 1) + ': ' + root_items[root_probs.index(max(root_probs))] + ' '
        self.m_patt_chart.addSeries(self.m_patt_series)
        self.m_posi_chart.addSeries(self.m_posi_series)
        self.m_reason_chart.addSeries(self.m_reason_series)
        self.m_root_chart.addSeries(self.m_root_series)
        self.m_patt_chart_2.addSeries(self.m_patt_series_2)
        self.m_posi_chart_2.addSeries(self.m_posi_series_2)
        self.m_reason_chart_2.addSeries(self.m_reason_series_2)
        self.m_root_chart_2.addSeries(self.m_root_series_2)
        # 填充表格
        self.ui.resultWidget_3.item(Fault.pattern_value.value).setText(whole_patt)
        self.ui.resultWidget_3.item(Fault.position_value.value).setText(whole_posi)
        self.ui.resultWidget_3.item(Fault.reason_value.value).setText(whole_reason)
        self.ui.resultWidget_3.item(Fault.root_value.value).setText(whole_root)
        # 填充表格
        self.ui.resultWidget_4.item(Fault.pattern_value.value).setText(subsys_patt)
        self.ui.resultWidget_4.item(Fault.position_value.value).setText(subsys_posi)
        self.ui.resultWidget_4.item(Fault.reason_value.value).setText(subsys_reason)
        self.ui.resultWidget_4.item(Fault.root_value.value).setText(subsys_root)

    @pyqtSlot()
    def on_loadButton_3_clicked(self):
        dev_1 = self.ui.devWidget.item(2).text()
        dev_2 = self.ui.devWidget_2.item(2).text()
        if dev_1.strip(' ') != '' and dev_2.strip(' ') == '':  # 只有设备1的数据
            dev_1 = int(dev_1)
            fault_data = self.read_fault_data(dev_1)
            self.plot_fault_bar([fault_data])
        elif dev_1.strip(' ') == '' and dev_2.strip(' ') != '':  # 只有设备2的数据
            dev_2 = int(dev_2)
            fault_data = self.read_fault_data(dev_2)
            self.plot_fault_bar([fault_data])
        elif dev_1.strip(' ') != '' and dev_2.strip(' ') != '':
            dev_1 = int(dev_1)
            dev_2 = int(dev_2)
            fault_data_1 = self.read_fault_data(dev_1)
            fault_data_2 = self.read_fault_data(dev_2)
            self.plot_fault_bar([fault_data_1, fault_data_2])
        # if not fault_data:
        #     msg_box = QMessageBox(QMessageBox.Warning,
        #                           '警告',
        #                           '请导入正确格式！',
        #                           QMessageBox.Ok)
        #     msg_box.exec_()
        #     return

    @pyqtSlot()
    def on_wholeButton_clicked(self):
        self.ui.faultWidget.setCurrentIndex(0)

    @pyqtSlot()
    def on_subSysButton_clicked(self):
        self.ui.faultWidget.setCurrentIndex(1)

    def save_view(self, view, save_path):
        '''
        抓取界面view并保存为图片文件
        :param view: view控件
        :param save_path: 保存路径
        :return:
        '''
        screen = QApplication.primaryScreen()
        pix = screen.grabWindow(view.winId())
        pix.save(save_path)

    def write_fault_result(self, filename):
        item_head = ['故障模式', '故障部位', '故障原因', '故障溯源']
        wb = Workbook()
        # 添加sheets
        whole_sheet = wb.create_sheet('整机', index=0)
        subsys_sheet = wb.create_sheet('子系统', index=1)
        # 先写入表头
        whole_sheet.append(item_head)
        subsys_sheet.append(item_head)
        # 写入结果
        whole_sheet.append([self.ui.resultWidget_3.item(2 * i + 1).text() for i in range(0, len(item_head))])
        subsys_sheet.append([self.ui.resultWidget_4.item(2 * i + 1).text() for i in range(0, len(item_head))])
        # 保存文档
        wb.save(filename)

    @pyqtSlot()
    def on_saveButton_3_clicked(self):
        save_dir = QFileDialog.getExistingDirectory(self, '选择文件夹', './')
        if save_dir != '':
            prefix = '故障分析_'
            whole = '整机_'
            subsys = '子系统_'
            # 保存整机分析图片
            if self.ui.faultWidget.currentIndex() == 0:
                self.save_view(self.ui.pattView, os.path.join(save_dir, '故障模式统计.jpg'))
                self.save_view(self.ui.posiView, os.path.join(save_dir, '故障部位统计.jpg'))
                self.save_view(self.ui.reasonView, os.path.join(save_dir, '故障原因统计.jpg'))
                self.save_view(self.ui.rootView, os.path.join(save_dir, '故障溯源统计.jpg'))
            else:
                # 保存子系统分析图片
                self.save_view(self.ui.pattView_2, os.path.join(save_dir, '故障模式统计.jpg'))
                self.save_view(self.ui.posiView_2, os.path.join(save_dir, '故障部位统计.jpg'))
                self.save_view(self.ui.reasonView_2, os.path.join(save_dir, '故障原因统计.jpg'))
                self.save_view(self.ui.rootView_2, os.path.join(save_dir, '故障溯源统计.jpg'))
            # 将结果写入表格
            self.write_fault_result(os.path.join(save_dir, '故障分析结果.xls'))
            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '保存成功',
                                  QMessageBox.Ok)
            msg_box.exec_()

    @pyqtSlot()
    def on_outputDataButton_3_clicked(self):
        item_head = ['序号', '故障模式', '故障部位', '故障原因', '故障溯源', '设备名称', '记录时间']
        wb = Workbook()
        # 添加sheets
        whole_sheet = wb.create_sheet('整机', index=0)
        subsys_sheet = wb.create_sheet('子系统', index=1)
        # 先写入表头
        whole_sheet.append(item_head)
        subsys_sheet.append(item_head)
        sql_str = 'select b.id, b.pattern, b.position, b.reason, b.root, b.status, ' \
                  '(select d.name from device d where d.id = b.dev_id), b.record_time from breakdown b'
        self.query.prepare(sql_str)
        if not self.query.exec_():
            print(self.query.lastError().text())
        else:
            while self.query.next():
                # 写入结果
                status = self.query.value(5)
                if 1 == status:
                    # 整机
                    item = [self.query.value(0), self.query.value(1),
                            self.query.value(2), self.query.value(3),
                            self.query.value(4), self.query.value(6), self.query.value(7)]
                    whole_sheet.append(item)
                else:
                    item = [self.query.value(0), self.query.value(1),
                            self.query.value(2), self.query.value(3),
                            self.query.value(4), self.query.value(6), self.query.value(7)]
                    subsys_sheet.append(item)
        filename, _ = QFileDialog.getSaveFileName(self, '保存文件', './', 'Excel 2007 (*.xlsx);;Excel 2003 (*.xls)')
        if filename.strip(' ') != '':
            wb.save(filename)
            msg_box = QMessageBox(QMessageBox.Information,
                                  '提示',
                                  '保存成功',
                                  QMessageBox.Ok)
            msg_box.exec_()

    @pyqtSlot()
    def on_manaFtRecButton_clicked(self):
        mt_widget = FaultListDialog()
        mt_widget.show()
        mt_widget.exec_()
        mt_widget.destroy()

    @pyqtSlot()
    def on_nextResultButton_clicked(self):
        page_num = self.ui.reliaResultWidget.count()
        index = self.ui.reliaResultWidget.currentIndex()
        index += 1
        if index == page_num:
            index = 0
        self.ui.reliaResultWidget.setCurrentIndex(index)

    def on_comboBox_currentIndexChanged(self):
        def check_not_nan(lambda_hat, beta_hat):
            flag = False
            if lambda_hat.strip(' ') != '' and beta_hat.strip(' ') != '':
                flag = True
            return flag

        def switch_algorithm(lambda_hat, beta_hat, lambda_map, beta_map, method, m_pdf_chart, m_pdf_series_ls,
                             m_pdf_series_map, m_cdf_chart, m_cdf_series_ls, m_cdf_series_map, m_relia_chart,
                             m_relia_series_ls, m_relia_series_map, m_fail_chart, m_fail_series_ls, m_fail_series_map,
                             resultWidget):
            if method == 0:
                m_pdf_chart.addSeries(m_pdf_series_ls)
                m_cdf_chart.addSeries(m_cdf_series_ls)
                m_relia_chart.addSeries(m_relia_series_ls)
                m_fail_chart.addSeries(m_fail_series_ls)

                m_pdf_chart.removeSeries(m_pdf_series_map)
                m_cdf_chart.removeSeries(m_cdf_series_map)
                m_relia_chart.removeSeries(m_relia_series_map)
                m_fail_chart.removeSeries(m_fail_series_map)

                # 填充表格
                resultWidget.item(Result.lamda_value.value).setText('%.4f' % lambda_hat)
                resultWidget.item(Result.beta_value.value).setText('%.4f' % beta_hat)
                resultWidget.item(Result.mtbf_value.value).setText('%.4f' % rf.mtbf(lambda_hat, beta_hat))
            else:
                m_pdf_chart.addSeries(m_pdf_series_map)
                m_cdf_chart.addSeries(m_cdf_series_map)
                m_relia_chart.addSeries(m_relia_series_map)
                m_fail_chart.addSeries(m_fail_series_map)

                m_pdf_chart.removeSeries(m_pdf_series_ls)
                m_cdf_chart.removeSeries(m_cdf_series_ls)
                m_relia_chart.removeSeries(m_relia_series_ls)
                m_fail_chart.removeSeries(m_fail_series_ls)
                # 填充表格
                resultWidget.item(Result.lamda_value.value).setText('%.4f' % lambda_map)
                resultWidget.item(Result.beta_value.value).setText('%.4f' % beta_map)
                resultWidget.item(Result.mtbf_value.value).setText('%.4f' % rf.mtbf(lambda_map, beta_map))

        method = self.ui.comboBox.currentIndex()  # 获取算法选择索引,0---最小二乘法, 1---贝叶斯估计法
        lambda_hat_1, beta_hat_1 = self.ui.dev1ResultWidget.item(
            Result.lamda_value.value).text(), self.ui.dev1ResultWidget.item(Result.beta_value.value).text()
        lambda_hat_2, beta_hat_2 = self.ui.dev2ResultWidget.item(
            Result.lamda_value.value).text(), self.ui.dev2ResultWidget.item(Result.beta_value.value).text()

        if check_not_nan(lambda_hat_1, beta_hat_1) and (not check_not_nan(lambda_hat_2, beta_hat_2)):
            switch_algorithm(self.lambda_hat_1, self.beta_hat_1, self.lambda_map_1, self.beta_map_1, method,
                             self.m_pdf_chart, self.m_pdf_series_ls, self.m_pdf_series_map, self.m_cdf_chart,
                             self.m_cdf_series_ls, self.m_cdf_series_map, self.m_relia_chart, self.m_relia_series_ls,
                             self.m_relia_series_map, self.m_fail_chart, self.m_fail_series_ls,
                             self.m_fail_series_map, self.ui.dev1ResultWidget)
        elif (not check_not_nan(lambda_hat_1, beta_hat_1)) and check_not_nan(lambda_hat_2, beta_hat_2):
            switch_algorithm(self.lambda_hat_2, self.beta_hat_2, self.lambda_map_2, self.beta_map_2, method,
                             self.m_pdf_chart, self.m_pdf_series_ls_2, self.m_pdf_series_map_2, self.m_cdf_chart,
                             self.m_cdf_series_ls_2, self.m_cdf_series_map_2, self.m_relia_chart,
                             self.m_relia_series_ls_2,
                             self.m_relia_series_map_2, self.m_fail_chart, self.m_fail_series_ls_2,
                             self.m_fail_series_map_2, self.ui.dev2ResultWidget)
        elif check_not_nan(lambda_hat_1, beta_hat_1) and check_not_nan(lambda_hat_2, beta_hat_2):
            switch_algorithm(self.lambda_hat_1, self.beta_hat_1, self.lambda_map_1, self.beta_map_1, method,
                             self.m_pdf_chart, self.m_pdf_series_ls, self.m_pdf_series_map, self.m_cdf_chart,
                             self.m_cdf_series_ls, self.m_cdf_series_map, self.m_relia_chart, self.m_relia_series_ls,
                             self.m_relia_series_map, self.m_fail_chart, self.m_fail_series_ls,
                             self.m_fail_series_map, self.ui.dev1ResultWidget)
            switch_algorithm(self.lambda_hat_2, self.beta_hat_2, self.lambda_map_2, self.beta_map_2, method,
                             self.m_pdf_chart, self.m_pdf_series_ls_2, self.m_pdf_series_map_2, self.m_cdf_chart,
                             self.m_cdf_series_ls_2, self.m_cdf_series_map_2, self.m_relia_chart,
                             self.m_relia_series_ls_2,
                             self.m_relia_series_map_2, self.m_fail_chart, self.m_fail_series_ls_2,
                             self.m_fail_series_map_2, self.ui.dev2ResultWidget)
        else:
            pass

    def on_comboBox_2_currentIndexChanged(self):
        def check_not_nan(lambda_hat, beta_hat):
            flag = False
            if lambda_hat.strip(' ') != '' and beta_hat.strip(' ') != '':
                flag = True
            return flag

        def switch_algorithm(lambda_hat, beta_hat, lambda_map, beta_map, method, m_pdf_chart, m_pdf_series_ls,
                             m_pdf_series_map, m_cdf_chart, m_cdf_series_ls, m_cdf_series_map, m_relia_chart,
                             m_relia_series_ls, m_relia_series_map, m_fail_chart, m_fail_series_ls, m_fail_series_map,
                             resultWidget):
            if method == 0:
                m_pdf_chart.addSeries(m_pdf_series_ls)
                m_cdf_chart.addSeries(m_cdf_series_ls)
                m_relia_chart.addSeries(m_relia_series_ls)
                m_fail_chart.addSeries(m_fail_series_ls)

                m_pdf_chart.removeSeries(m_pdf_series_map)
                m_cdf_chart.removeSeries(m_cdf_series_map)
                m_relia_chart.removeSeries(m_relia_series_map)
                m_fail_chart.removeSeries(m_fail_series_map)

                # 填充表格
                resultWidget.item(Result.lamda_value.value).setText('%.4f' % lambda_hat)
                resultWidget.item(Result.beta_value.value).setText('%.4f' % beta_hat)
                resultWidget.item(Result.mtbf_value.value).setText('%.4f' % rf.mtbf(lambda_hat, beta_hat))
            else:
                m_pdf_chart.addSeries(m_pdf_series_map)
                m_cdf_chart.addSeries(m_cdf_series_map)
                m_relia_chart.addSeries(m_relia_series_map)
                m_fail_chart.addSeries(m_fail_series_map)

                m_pdf_chart.removeSeries(m_pdf_series_ls)
                m_cdf_chart.removeSeries(m_cdf_series_ls)
                m_relia_chart.removeSeries(m_relia_series_ls)
                m_fail_chart.removeSeries(m_fail_series_ls)
                # 填充表格
                resultWidget.item(Result.lamda_value.value).setText('%.4f' % lambda_map)
                resultWidget.item(Result.beta_value.value).setText('%.4f' % beta_map)
                resultWidget.item(Result.mtbf_value.value).setText('%.4f' % rf.mtbf(lambda_map, beta_map))

        method = self.ui.comboBox_2.currentIndex()  # 获取算法选择索引,0---最小二乘法, 1---贝叶斯估计法
        lambda_hat_mt_1, beta_hat_mt_1 = self.ui.dev1MtResultWidget.item(
            Result.lamda_value.value).text(), self.ui.dev1MtResultWidget.item(Result.beta_value.value).text()
        lambda_hat_mt_2, beta_hat_mt_2 = self.ui.dev2MtResultWidget.item(
            Result.lamda_value.value).text(), self.ui.dev2MtResultWidget.item(Result.beta_value.value).text()

        if check_not_nan(lambda_hat_mt_1, beta_hat_mt_1) and (not check_not_nan(lambda_hat_mt_2, beta_hat_mt_2)):
            switch_algorithm(self.lambda_hat_mt_1, self.beta_hat_mt_1, self.lambda_map_mt_1, self.beta_map_mt_1, method,
                             self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls, self.m_pdf_series_maintain_map,
                             self.m_cdf_chart_maintain,
                             self.m_cdf_series_maintain_ls, self.m_cdf_series_maintain_map, self.m_relia_chart_maintain,
                             self.m_relia_series_maintain_ls,
                             self.m_relia_series_maintain_map, self.m_fail_chart_maintain,
                             self.m_fail_series_maintain_ls,
                             self.m_fail_series_maintain_map, self.ui.dev1MtResultWidget)
        elif (not check_not_nan(lambda_hat_mt_1, beta_hat_mt_1)) and check_not_nan(lambda_hat_mt_2, beta_hat_mt_2):
            switch_algorithm(self.lambda_hat_mt_2, self.beta_hat_mt_2, self.lambda_map_mt_2, self.beta_map_mt_2, method,
                             self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls_2,
                             self.m_pdf_series_maintain_map_2, self.m_cdf_chart_maintain,
                             self.m_cdf_series_maintain_ls_2, self.m_cdf_series_maintain_map_2,
                             self.m_relia_chart_maintain,
                             self.m_relia_series_maintain_ls_2,
                             self.m_relia_series_maintain_map_2, self.m_fail_chart_maintain,
                             self.m_fail_series_maintain_ls_2,
                             self.m_fail_series_maintain_map_2, self.ui.dev2MtResultWidget)
        elif check_not_nan(lambda_hat_mt_1, beta_hat_mt_1) and check_not_nan(lambda_hat_mt_2, beta_hat_mt_2):
            switch_algorithm(self.lambda_hat_mt_1, self.beta_hat_mt_1, self.lambda_map_mt_1, self.beta_map_mt_1, method,
                             self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls, self.m_pdf_series_maintain_map,
                             self.m_cdf_chart_maintain,
                             self.m_cdf_series_maintain_ls, self.m_cdf_series_maintain_map, self.m_relia_chart_maintain,
                             self.m_relia_series_maintain_ls,
                             self.m_relia_series_maintain_map, self.m_fail_chart_maintain,
                             self.m_fail_series_maintain_ls,
                             self.m_fail_series_maintain_map, self.ui.dev1MtResultWidget)
            switch_algorithm(self.lambda_hat_mt_2, self.beta_hat_mt_2, self.lambda_map_mt_2, self.beta_map_mt_2, method,
                             self.m_pdf_chart_maintain, self.m_pdf_series_maintain_ls_2,
                             self.m_pdf_series_maintain_map_2, self.m_cdf_chart_maintain,
                             self.m_cdf_series_maintain_ls_2, self.m_cdf_series_maintain_map_2,
                             self.m_relia_chart_maintain,
                             self.m_relia_series_maintain_ls_2,
                             self.m_relia_series_maintain_map_2, self.m_fail_chart_maintain,
                             self.m_fail_series_maintain_ls_2,
                             self.m_fail_series_maintain_map_2, self.ui.dev2MtResultWidget)
        else:
            pass

    def on_axisComboBox_currentIndexChanged(self):
        def check_data_exist(data_table):
            if len(data_table) <= 0 or (
                    data_table['x_table'] == [[], [], [], []] and data_table['y_table'] == [[], [], [], []] and
                    data_table['z_table'] == [[], [], [], []]):
                flag = False
            else:
                flag = True
            return flag

        dev_1 = self.ui.devWidget.item(2).text()
        dev_2 = self.ui.devWidget_2.item(2).text()
        if dev_1.strip(' ') != '' and dev_2.strip(' ') == '':
            if not check_data_exist(self.data_table_of_dev_1):
                msg_box = QMessageBox(QMessageBox.Warning,
                                      '提示',
                                      '数据库无记录！',
                                      QMessageBox.Ok)
                msg_box.exec_()
                return
            else:
                # ---原始数据可视化---
                axis = self.ui.axisComboBox.currentIndex()  # 0---x axis, 1---y axis, 2---z axis
                if axis == 0:
                    self.plot_raw_data(self.data_table_of_dev_1['x_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                elif axis == 1:
                    self.plot_raw_data(self.data_table_of_dev_1['y_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                else:
                    self.plot_raw_data(self.data_table_of_dev_1['z_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
        elif dev_1.strip(' ') == '' and dev_2.strip(' ') != '':
            if not check_data_exist(self.data_table_of_dev_2):
                msg_box = QMessageBox(QMessageBox.Warning,
                                      '提示',
                                      '数据库无记录！',
                                      QMessageBox.Ok)
                msg_box.exec_()
                return
            else:
                # ---原始数据可视化---
                axis = self.ui.axisComboBox.currentIndex()  # 0---x axis, 1---y axis, 2---z axis
                if axis == 0:
                    self.plot_raw_data(self.data_table_of_dev_2['x_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                elif axis == 1:
                    self.plot_raw_data(self.data_table_of_dev_2['y_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                else:
                    self.plot_raw_data(self.data_table_of_dev_2['z_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
        elif dev_1.strip(' ') != '' and dev_2.strip(' ') != '':
            if not (check_data_exist(self.data_table_of_dev_1) or check_data_exist(self.data_table_of_dev_2)):
                msg_box = QMessageBox(QMessageBox.Warning,
                                      '提示',
                                      '数据库无记录！',
                                      QMessageBox.Ok)
                msg_box.exec_()
                return
            else:
                # ---原始数据可视化---
                axis = self.ui.axisComboBox.currentIndex()  # 0---x axis, 1---y axis, 2---z axis
                if axis == 0:
                    self.plot_raw_data(self.data_table_of_dev_1['x_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                    self.plot_raw_data(self.data_table_of_dev_2['x_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                elif axis == 1:
                    self.plot_raw_data(self.data_table_of_dev_1['y_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                    self.plot_raw_data(self.data_table_of_dev_2['y_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
                else:
                    self.plot_raw_data(self.data_table_of_dev_1['z_table'], self.m_env_temp_series,
                                       self.m_env_temp_scatter, self.m_kni_temp_series, self.m_kni_temp_scatter,
                                       self.m_rpa_series, self.m_rpa_scatter, self.m_stra_series, self.m_stra_scatter)
                    self.plot_raw_data(self.data_table_of_dev_2['z_table'], self.m_env_temp_series_2,
                                       self.m_env_temp_scatter_2, self.m_kni_temp_series_2, self.m_kni_temp_scatter_2,
                                       self.m_rpa_series_2, self.m_rpa_scatter_2, self.m_stra_series_2,
                                       self.m_stra_scatter_2)
        else:
            pass

    def on_devComboBox_currentIndexChanged(self):
        dev_name = self.ui.devComboBox.currentText()
        if dev_name != '选择设备':
            query_sql = 'select id, num, name from device where name = :name'
            self.query.prepare(query_sql)
            self.query.bindValue(':name', dev_name)
            if not self.query.exec_():
                print(self.query.lastError())
            else:
                while self.query.next():
                    id = self.query.value(0)
                    num = self.query.value(1)
                    name = self.query.value(2)
                    self.ui.devWidget.item(2).setText(str(id))
                    self.ui.devWidget.item(4).setText(str(num))
                    self.ui.devWidget.item(6).setText(str(name))
            # 不能选择同一个设备
            dev2_name = self.ui.devComboBox_2.currentText()
            if dev_name == dev2_name:
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '请选择不同的设备！',
                                      QMessageBox.Ok)
                msg_box.exec_()
                self.ui.devComboBox.setCurrentIndex(0)
                self.ui.devWidget.item(2).setText('')
                self.ui.devWidget.item(4).setText('')
                self.ui.devWidget.item(6).setText('')

    def on_devComboBox_2_currentIndexChanged(self):
        dev2_name = self.ui.devComboBox_2.currentText()
        if dev2_name != '选择设备':
            query_sql = 'select id, num, name from device where name = :name'
            self.query.prepare(query_sql)
            self.query.bindValue(':name', dev2_name)
            if not self.query.exec_():
                print(self.query.lastError())
            else:
                while self.query.next():
                    id = self.query.value(0)
                    num = self.query.value(1)
                    name = self.query.value(2)
                    self.ui.devWidget_2.item(2).setText(str(id))
                    self.ui.devWidget_2.item(4).setText(str(num))
                    self.ui.devWidget_2.item(6).setText(str(name))
            # 不能选择同一个设备
            dev_name = self.ui.devComboBox.currentText()
            if dev_name == dev2_name:
                msg_box = QMessageBox(QMessageBox.Information,
                                      '提示',
                                      '请选择不同的设备！',
                                      QMessageBox.Ok)
                msg_box.exec_()
                self.ui.devComboBox_2.setCurrentIndex(0)
                self.ui.devWidget_2.item(2).setText('')
                self.ui.devWidget_2.item(4).setText('')
                self.ui.devWidget_2.item(6).setText('')

    def on_m_pdf_series_hovered(self, point, state):
        sender = self.sender()
        if sender == self.m_pdf_series_ls or sender == self.m_pdf_series_map:
            resultWidget = self.ui.dev1ResultWidget
        else:
            resultWidget = self.ui.dev2ResultWidget
        if state:
            curr_time = point.x()
            pdf = point.y()
            lambda_hat, beta_hat = float(resultWidget.item(Result.lamda_value.value).text()), float(
                resultWidget.item(Result.beta_value.value).text())
            cdf = rf.Weibull_cdf(lambda_hat, beta_hat, curr_time)
            relia = rf.reliability(lambda_hat, beta_hat, curr_time)
            fali = rf.failure_rate(lambda_hat, beta_hat, curr_time)

            resultWidget.item(Result.curr_time_value.value).setText('%f' % curr_time)
            resultWidget.item(Result.pdf_value.value).setText('%f' % pdf)
            resultWidget.item(Result.cdf_value.value).setText('%f' % cdf)
            resultWidget.item(Result.relia_value.value).setText('%f' % relia)
            resultWidget.item(Result.fali_value.value).setText('%f' % fali)

    def on_m_cdf_series_hovered(self, point, state):
        sender = self.sender()
        if sender == self.m_cdf_series_ls or sender == self.m_cdf_series_map:
            resultWidget = self.ui.dev1ResultWidget
        else:
            resultWidget = self.ui.dev2ResultWidget
        if state:
            curr_time = point.x()
            cdf = point.y()
            lambda_hat, beta_hat = float(resultWidget.item(Result.lamda_value.value).text()), float(
                resultWidget.item(Result.beta_value.value).text())
            pdf = rf.Weibull_pdf(lambda_hat, beta_hat, curr_time)
            relia = rf.reliability(lambda_hat, beta_hat, curr_time)
            fali = rf.failure_rate(lambda_hat, beta_hat, curr_time)

            resultWidget.item(Result.curr_time_value.value).setText('%f' % curr_time)
            resultWidget.item(Result.pdf_value.value).setText('%f' % pdf)
            resultWidget.item(Result.cdf_value.value).setText('%f' % cdf)
            resultWidget.item(Result.relia_value.value).setText('%f' % relia)
            resultWidget.item(Result.fali_value.value).setText('%f' % fali)

    def on_m_relia_series_hovered(self, point, state):
        sender = self.sender()
        if sender == self.m_relia_series_ls or sender == self.m_relia_series_map:
            resultWidget = self.ui.dev1ResultWidget
        else:
            resultWidget = self.ui.dev2ResultWidget
        if state:
            curr_time = point.x()
            relia = point.y()
            lambda_hat, beta_hat = float(resultWidget.item(Result.lamda_value.value).text()), float(
                resultWidget.item(Result.beta_value.value).text())
            cdf = rf.Weibull_cdf(lambda_hat, beta_hat, curr_time)
            pdf = rf.Weibull_pdf(lambda_hat, beta_hat, curr_time)
            fali = rf.failure_rate(lambda_hat, beta_hat, curr_time)
            resultWidget.item(Result.curr_time_value.value).setText('%f' % curr_time)
            resultWidget.item(Result.pdf_value.value).setText('%f' % pdf)
            resultWidget.item(Result.cdf_value.value).setText('%f' % cdf)
            resultWidget.item(Result.relia_value.value).setText('%f' % relia)
            resultWidget.item(Result.fali_value.value).setText('%f' % fali)

    def on_m_fali_series_hovered(self, point, state):
        sender = self.sender()
        if sender == self.m_fail_series_ls or sender == self.m_fail_series_map:
            resultWidget = self.ui.dev1ResultWidget
        else:
            resultWidget = self.ui.dev2ResultWidget
        if state:
            curr_time = point.x()
            fali = point.y()
            lambda_hat, beta_hat = float(resultWidget.item(Result.lamda_value.value).text()), float(
                resultWidget.item(Result.beta_value.value).text())
            cdf = rf.Weibull_cdf(lambda_hat, beta_hat, curr_time)
            relia = rf.reliability(lambda_hat, beta_hat, curr_time)
            pdf = rf.Weibull_pdf(lambda_hat, beta_hat, curr_time)
            resultWidget.item(Result.curr_time_value.value).setText('%f' % curr_time)
            resultWidget.item(Result.pdf_value.value).setText('%f' % pdf)
            resultWidget.item(Result.cdf_value.value).setText('%f' % cdf)
            resultWidget.item(Result.relia_value.value).setText('%f' % relia)
            resultWidget.item(Result.fali_value.value).setText('%f' % fali)

    def on_m_pdf_series_maintain_hovered(self, point, state):
        sender = self.sender()
        if sender == self.m_pdf_series_maintain_ls or sender == self.m_pdf_series_maintain_map:
            resultWidget = self.ui.dev1MtResultWidget
        else:
            resultWidget = self.ui.dev2MtResultWidget
        if state:
            curr_time = point.x()
            pdf = point.y()
            lambda_hat, beta_hat = float(resultWidget.item(Result.lamda_value.value).text()), float(
                resultWidget.item(Result.beta_value.value).text())
            cdf = rf.Weibull_cdf(lambda_hat, beta_hat, curr_time)
            relia = rf.reliability(lambda_hat, beta_hat, curr_time)
            fali = rf.failure_rate(lambda_hat, beta_hat, curr_time)
            resultWidget.item(Result.curr_time_value.value).setText('%f' % curr_time)
            resultWidget.item(Result.pdf_value.value).setText('%f' % pdf)
            resultWidget.item(Result.cdf_value.value).setText('%f' % cdf)
            resultWidget.item(Result.relia_value.value).setText('%f' % relia)
            resultWidget.item(Result.fali_value.value).setText('%f' % fali)

    def on_m_cdf_series_maintain_hovered(self, point, state):
        sender = self.sender()
        if sender == self.m_cdf_series_maintain_ls or sender == self.m_cdf_series_maintain_map:
            resultWidget = self.ui.dev1MtResultWidget
        else:
            resultWidget = self.ui.dev2MtResultWidget
        if state:
            curr_time = point.x()
            cdf = point.y()
            lambda_hat, beta_hat = float(resultWidget.item(Result.lamda_value.value).text()), float(
                resultWidget.item(Result.beta_value.value).text())
            pdf = rf.Weibull_pdf(lambda_hat, beta_hat, curr_time)
            relia = rf.reliability(lambda_hat, beta_hat, curr_time)
            fali = rf.failure_rate(lambda_hat, beta_hat, curr_time)
            resultWidget.item(Result.curr_time_value.value).setText('%f' % curr_time)
            resultWidget.item(Result.pdf_value.value).setText('%f' % pdf)
            resultWidget.item(Result.cdf_value.value).setText('%f' % cdf)
            resultWidget.item(Result.relia_value.value).setText('%f' % relia)
            resultWidget.item(Result.fali_value.value).setText('%f' % fali)

    def on_m_relia_series_maintain_hovered(self, point, state):
        sender = self.sender()
        if sender == self.m_relia_series_maintain_ls or sender == self.m_relia_series_maintain_map:
            resultWidget = self.ui.dev1MtResultWidget
        else:
            resultWidget = self.ui.dev2MtResultWidget
        if state:
            curr_time = point.x()
            relia = point.y()
            lambda_hat, beta_hat = float(resultWidget.item(Result.lamda_value.value).text()), float(
                resultWidget.item(Result.beta_value.value).text())
            cdf = rf.Weibull_cdf(lambda_hat, beta_hat, curr_time)
            pdf = rf.Weibull_pdf(lambda_hat, beta_hat, curr_time)
            fali = rf.failure_rate(lambda_hat, beta_hat, curr_time)
            resultWidget.item(Result.curr_time_value.value).setText('%f' % curr_time)
            resultWidget.item(Result.pdf_value.value).setText('%f' % pdf)
            resultWidget.item(Result.cdf_value.value).setText('%f' % cdf)
            resultWidget.item(Result.relia_value.value).setText('%f' % relia)
            resultWidget.item(Result.fali_value.value).setText('%f' % fali)

    def on_m_fali_series_maintain_hovered(self, point, state):
        sender = self.sender()
        if sender == self.m_fail_series_maintain_ls or sender == self.m_fail_series_maintain_map:
            resultWidget = self.ui.dev1MtResultWidget
        else:
            resultWidget = self.ui.dev2MtResultWidget
        if state:
            curr_time = point.x()
            fali = point.y()
            lambda_hat, beta_hat = float(resultWidget.item(Result.lamda_value.value).text()), float(
                resultWidget.item(Result.beta_value.value).text())
            cdf = rf.Weibull_cdf(lambda_hat, beta_hat, curr_time)
            relia = rf.reliability(lambda_hat, beta_hat, curr_time)
            pdf = rf.Weibull_pdf(lambda_hat, beta_hat, curr_time)
            resultWidget.item(Result.curr_time_value.value).setText('%f' % curr_time)
            resultWidget.item(Result.pdf_value.value).setText('%f' % pdf)
            resultWidget.item(Result.cdf_value.value).setText('%f' % cdf)
            resultWidget.item(Result.relia_value.value).setText('%f' % relia)
            resultWidget.item(Result.fali_value.value).setText('%f' % fali)


if __name__ == '__main__':
    import sys

    app = QApplication(sys.argv)
    ui = MainWindow()
    ui.showMaximized()
    sys.exit(app.exec_())
